"""Local resume document import and accessible DOCX export helpers."""

from __future__ import annotations

import io
import math
import re
import zipfile
from dataclasses import dataclass
from typing import Callable, Iterable

import pymupdf
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


MAX_DOCUMENT_PAGES = 12
MAX_DOCX_EXPANDED_BYTES = 50 * 1024 * 1024
MAX_DOCX_PARTS = 1_000
MAX_IMPORTED_RESUME_CHARACTERS = 200_000
MAX_OCR_RENDER_DIMENSION = 2_500
MAX_OCR_RENDER_PIXELS = 5_000_000
MAX_OCR_IMAGE_BYTES = 12 * 1024 * 1024
MIN_READABLE_PAGE_CHARACTERS = 20
OCR_RENDER_SCALE = 1.5
SUPPORTED_RESUME_EXTENSIONS = frozenset({".txt", ".md", ".docx", ".pdf"})


@dataclass(frozen=True)
class ImportedResume:
    text: str
    source_format: str
    ocr_used: bool = False
    page_count: int | None = None


class ResumeDocumentError(ValueError):
    """A bounded, user-safe document import failure."""

    def __init__(self, message: str, status_code: int = 422):
        super().__init__(message)
        self.status_code = status_code


def _checkpoint(callback: Callable[[], None] | None) -> None:
    if callback:
        callback()


def _normalize_text(value: str) -> str:
    lines = [re.sub(r"[\t ]+", " ", line).rstrip() for line in (value or "").splitlines()]
    normalized: list[str] = []
    blank = False
    for line in lines:
        if line.strip():
            normalized.append(line.strip())
            blank = False
        elif normalized and not blank:
            normalized.append("")
            blank = True
    return "\n".join(normalized).strip()


def _validate_text_length(value: str) -> None:
    if len(value) > MAX_IMPORTED_RESUME_CHARACTERS:
        raise ResumeDocumentError(
            f"Imported resume text may contain at most {MAX_IMPORTED_RESUME_CHARACTERS:,} characters."
        )


def _validate_docx_container(content: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            members = archive.infolist()
            if len(members) > MAX_DOCX_PARTS:
                raise ResumeDocumentError("The DOCX contains too many embedded parts to import safely.")
            if any(member.flag_bits & 0x1 for member in members):
                raise ResumeDocumentError("Password-protected DOCX files cannot be imported.")
            if sum(member.file_size for member in members) > MAX_DOCX_EXPANDED_BYTES:
                raise ResumeDocumentError("The expanded DOCX is too large to import safely.")
            names = {member.filename for member in members}
            if "word/document.xml" not in names or "[Content_Types].xml" not in names:
                raise ResumeDocumentError("The uploaded file is not a readable DOCX document.")
    except ResumeDocumentError:
        raise
    except (zipfile.BadZipFile, OSError) as exc:
        raise ResumeDocumentError("The uploaded DOCX is damaged or not a valid Word document.") from exc


def _iter_document_blocks(document: DocumentObject) -> Iterable[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _paragraph_markdown(paragraph: Paragraph) -> str:
    text = _normalize_text(paragraph.text)
    if not text:
        return ""
    style_name = (paragraph.style.name if paragraph.style else "").strip().lower()
    style_id = (paragraph.style.style_id if paragraph.style else "").strip().lower()
    heading_match = re.fullmatch(r"heading\s*([1-9])", style_name) or re.fullmatch(r"heading([1-9])", style_id)
    if heading_match:
        level = min(int(heading_match.group(1)), 3)
        return f"{'#' * level} {text}"

    properties = paragraph._p.pPr
    has_numbering = properties is not None and properties.numPr is not None
    if style_name.startswith("list bullet") or style_id.startswith("listbullet"):
        return f"- {text}"
    if style_name.startswith("list number") or style_id.startswith("listnumber"):
        return f"1. {text}"
    if has_numbering:
        return f"- {text}"
    if re.match(r"^[-*•]\s+", text):
        return re.sub(r"^[-*•]\s+", "- ", text)
    if re.match(r"^\d+[.)]\s+", text):
        return re.sub(r"^\d+[.)]\s+", "1. ", text)
    return text


def _table_markdown(table: Table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = [_normalize_text(cell.text) for cell in row.cells]
        values = [value for index, value in enumerate(cells) if value and value not in cells[:index]]
        if values:
            rows.append(" | ".join(values))
    return "\n".join(rows)


def import_docx(content: bytes, checkpoint: Callable[[], None] | None = None) -> ImportedResume:
    _checkpoint(checkpoint)
    _validate_docx_container(content)
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise ResumeDocumentError("The uploaded DOCX could not be read.") from exc

    blocks: list[str] = []
    for block in _iter_document_blocks(document):
        _checkpoint(checkpoint)
        text = _paragraph_markdown(block) if isinstance(block, Paragraph) else _table_markdown(block)
        if text:
            blocks.append(text)
            _validate_text_length("\n\n".join(blocks))
    result = _normalize_text("\n\n".join(blocks))
    if not result:
        raise ResumeDocumentError("The DOCX did not contain readable resume text.")
    return ImportedResume(text=result, source_format="docx")


def _page_has_readable_text(text: str) -> bool:
    return len(re.sub(r"\W", "", text or "", flags=re.UNICODE)) >= MIN_READABLE_PAGE_CHARACTERS


def import_pdf(
    content: bytes,
    allow_ocr: bool = False,
    ocr_images: Callable[[list[bytes], Callable[[], None] | None], list[str]] | None = None,
    checkpoint: Callable[[], None] | None = None,
) -> ImportedResume:
    _checkpoint(checkpoint)
    try:
        document = pymupdf.open(stream=content, filetype="pdf")
    except Exception as exc:
        raise ResumeDocumentError("The uploaded PDF is damaged, encrypted, or unreadable.") from exc

    try:
        if document.needs_pass:
            raise ResumeDocumentError("Password-protected PDF files cannot be imported.")
        if document.page_count < 1:
            raise ResumeDocumentError("The PDF did not contain any pages.")
        if document.page_count > MAX_DOCUMENT_PAGES:
            raise ResumeDocumentError(f"Resume PDFs may contain at most {MAX_DOCUMENT_PAGES} pages.")

        page_text: list[str] = []
        scanned_indexes: list[int] = []
        for page_index in range(document.page_count):
            _checkpoint(checkpoint)
            extracted = _normalize_text(document.load_page(page_index).get_text("text"))
            _validate_text_length("\n\n".join((*page_text, extracted)))
            page_text.append(extracted)
            if not _page_has_readable_text(extracted):
                scanned_indexes.append(page_index)

        if scanned_indexes:
            if not allow_ocr:
                page_label = "page appears" if len(scanned_indexes) == 1 else "pages appear"
                raise ResumeDocumentError(
                    f"{len(scanned_indexes)} PDF {page_label} scanned or image-only. "
                    "Select the AI OCR option and import again if you consent to send those page images to your selected AI provider."
                )
            if ocr_images is None:
                raise ResumeDocumentError("AI OCR is not available for this import.")
            images: list[bytes] = []
            for page_index in scanned_indexes:
                _checkpoint(checkpoint)
                page = document.load_page(page_index)
                width = float(page.rect.width)
                height = float(page.rect.height)
                if width <= 0 or height <= 0 or not math.isfinite(width * height):
                    raise ResumeDocumentError(f"PDF page {page_index + 1} has invalid dimensions.")
                scale = min(
                    OCR_RENDER_SCALE,
                    MAX_OCR_RENDER_DIMENSION / max(width, height),
                    math.sqrt((MAX_OCR_RENDER_PIXELS * 0.98) / (width * height)),
                )
                if scale <= 0:
                    raise ResumeDocumentError(f"PDF page {page_index + 1} cannot be rendered safely for OCR.")
                pixmap = page.get_pixmap(
                    matrix=pymupdf.Matrix(scale, scale),
                    colorspace=pymupdf.csRGB,
                    alpha=False,
                )
                image = pixmap.tobytes("png")
                if len(image) > MAX_OCR_IMAGE_BYTES:
                    raise ResumeDocumentError(f"PDF page {page_index + 1} is too detailed to send safely for OCR.")
                images.append(image)
            ocr_text = ocr_images(images, checkpoint)
            if len(ocr_text) != len(scanned_indexes):
                raise ResumeDocumentError("The AI OCR provider returned an incomplete page set.")
            for page_index, text in zip(scanned_indexes, ocr_text):
                normalized = _normalize_text(text)
                if not normalized:
                    raise ResumeDocumentError(f"AI OCR could not read PDF page {page_index + 1}.")
                page_text[page_index] = normalized

        result = _normalize_text("\n\n".join(page_text))
        _validate_text_length(result)
        if not result:
            raise ResumeDocumentError("The PDF did not contain readable resume text.")
        return ImportedResume(
            text=result,
            source_format="pdf",
            ocr_used=bool(scanned_indexes),
            page_count=document.page_count,
        )
    finally:
        document.close()


def import_resume_document(
    filename: str,
    content: bytes,
    allow_ocr: bool = False,
    ocr_images: Callable[[list[bytes], Callable[[], None] | None], list[str]] | None = None,
    checkpoint: Callable[[], None] | None = None,
) -> ImportedResume:
    suffix = ("." + (filename or "").rsplit(".", 1)[-1].lower()) if "." in (filename or "") else ""
    if suffix not in SUPPORTED_RESUME_EXTENSIONS:
        raise ResumeDocumentError(
            "Resume imports support .txt, .md, .docx, and .pdf files.",
            status_code=400,
        )
    if suffix in {".txt", ".md"}:
        try:
            text = content.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise ResumeDocumentError("Text resume files must use UTF-8 encoding.", status_code=400) from exc
        normalized = _normalize_text(text)
        if not normalized:
            raise ResumeDocumentError("The resume file did not contain readable text.")
        _validate_text_length(normalized)
        return ImportedResume(text=normalized, source_format=suffix[1:])
    if suffix == ".docx":
        return import_docx(content, checkpoint)
    return import_pdf(content, allow_ocr, ocr_images, checkpoint)


_PHONE_TOKEN = r"(?<!\d)(?:\+?1[\s.-]?)?\(?\d{3}\)?[\s.-]\d{3}[\s.-]\d{4}(?!\d)"
_INLINE_TOKEN = re.compile(
    rf"(\[[^\]]+\]\(https?://[^\s)]+\)|https?://[^\s<>]+|[\w.+-]+@[\w.-]+\.[A-Za-z]{{2,}}|{_PHONE_TOKEN}|\*\*[^*]+\*\*|\*[^*]+\*)"
)


def _add_hyperlink(paragraph: Paragraph, text: str, target: str) -> None:
    relationship_id = paragraph.part.relate_to(
        target,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    run.append(properties)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_inline_content(paragraph: Paragraph, text: str) -> None:
    cursor = 0
    for match in _INLINE_TOKEN.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        token = match.group(0)
        markdown_link = re.fullmatch(r"\[([^\]]+)\]\((https?://[^\s)]+)\)", token)
        if markdown_link:
            _add_hyperlink(paragraph, markdown_link.group(1), markdown_link.group(2))
        elif token.startswith(("http://", "https://")):
            target = token.rstrip(".,;:")
            _add_hyperlink(paragraph, target, target)
            if len(target) != len(token):
                paragraph.add_run(token[len(target):])
        elif "@" in token:
            _add_hyperlink(paragraph, token, f"mailto:{token}")
        elif re.fullmatch(_PHONE_TOKEN, token):
            prefix = "+" if token.strip().startswith("+") else ""
            digits = re.sub("[^0-9]", "", token)
            _add_hyperlink(paragraph, token, f"tel:{prefix}{digits}")
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _set_default_language(document: DocumentObject, language: str = "en-US") -> None:
    styles = document.styles.element
    defaults = styles.find(qn("w:docDefaults"))
    if defaults is None:
        defaults = OxmlElement("w:docDefaults")
        styles.insert(0, defaults)
    run_default = defaults.find(qn("w:rPrDefault"))
    if run_default is None:
        run_default = OxmlElement("w:rPrDefault")
        defaults.append(run_default)
    run_properties = run_default.find(qn("w:rPr"))
    if run_properties is None:
        run_properties = OxmlElement("w:rPr")
        run_default.append(run_properties)
    language_node = run_properties.find(qn("w:lang"))
    if language_node is None:
        language_node = OxmlElement("w:lang")
        run_properties.append(language_node)
    language_node.set(qn("w:val"), language)
    language_node.set(qn("w:eastAsia"), language)


def _configure_accessible_styles(document: DocumentObject) -> None:
    accent = RGBColor(31, 78, 121)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, 18, 10),
        "Heading 2": (13, 14, 7),
        "Heading 3": (12, 10, 5),
    }
    for name, (size, before, after) in heading_tokens.items():
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = accent
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.1875)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.25

    if "Resume Contact" not in document.styles:
        contact = document.styles.add_style("Resume Contact", WD_STYLE_TYPE.PARAGRAPH)
        contact.base_style = normal
        contact.font.size = Pt(10)
        contact.paragraph_format.space_after = Pt(4)


def build_accessible_resume_docx(
    markdown_text: str,
    candidate_name: str = "",
    title: str = "Tailored Resume",
) -> bytes:
    content = (markdown_text or "").strip()
    if not content:
        raise ResumeDocumentError("Tailored resume text is unavailable for DOCX export.")

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    _set_default_language(document)
    _configure_accessible_styles(document)

    document.core_properties.title = title
    document.core_properties.subject = "Accessible, editable resume"
    document.core_properties.category = "Resume"
    if candidate_name.strip():
        document.core_properties.author = candidate_name.strip()

    paragraph_count = 0
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line or re.fullmatch(r"-{3,}", line):
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        bullet = re.match(r"^[-*•]\s+(.+)$", line)
        numbered = re.match(r"^\d+[.)]\s+(.+)$", line)
        if heading:
            paragraph = document.add_paragraph(style=f"Heading {len(heading.group(1))}")
            _add_inline_content(paragraph, heading.group(2))
        elif bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline_content(paragraph, bullet.group(1))
        elif numbered:
            paragraph = document.add_paragraph(style="List Number")
            _add_inline_content(paragraph, numbered.group(1))
        else:
            style = "Resume Contact" if paragraph_count <= 1 and re.search(r"https?://|@|\b\d{3}[-.) ]", line) else "Normal"
            paragraph = document.add_paragraph(style=style)
            _add_inline_content(paragraph, line)
        paragraph.paragraph_format.widow_control = True
        paragraph_count += 1

    if not paragraph_count:
        raise ResumeDocumentError("Tailored resume text is unavailable for DOCX export.")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
