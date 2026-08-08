import json
import io
import os
import re
import sqlite3
import tempfile
import unittest
import zipfile
from datetime import date
from pathlib import Path
from unittest.mock import patch
from uuid import uuid4

import app as app_module
import ai_providers as ai_providers_module
import maps_providers as maps_providers_module
import database as database_module
import materials as materials_module
import operations as operations_module
import searcher as searcher_module
import source_diagnostics as source_diagnostics_module
from application_insights import build_application_insights
import pymupdf
from docx import Document as WordDocument
from fastapi.testclient import TestClient
from ai_providers import AIProviderError, AIProviderSettings, CapabilityResponse, extract_text_from_images, generate_structured, settings_from_profile
from lifecycle import undo_latest_lifecycle_change, update_lifecycle
from materials import material_download_name, persist_cover_letter, resolve_output_file
from resume_documents import (
    MAX_IMPORTED_RESUME_CHARACTERS,
    MAX_OCR_RENDER_DIMENSION,
    MAX_OCR_RENDER_PIXELS,
    ResumeDocumentError,
    build_accessible_resume_docx,
    import_resume_document,
)
from maps_providers import (
    HeadquartersResult,
    MapsProviderSettings,
    lookup_headquarters,
    maps_provider_ready,
    resolve_headquarters,
)
from database import get_db_connection
from dependency_lock import (
    dependency_fingerprint,
    environment_matches_lock,
    locked_requirements,
    stamp_is_current,
    write_stamp,
)
from job_cleanup import apply_cleanup, cleanup_preview
from job_filters import derive_job_filter_facets
from job_suppressions import job_url_fingerprint, record_job_suppression
from base_resumes import (
    LastBaseResumeError,
    activate_base_resume,
    delete_base_resume,
    get_base_resume,
    list_versions,
    professional_evidence_markdown,
    restore_version,
    save_base_resume,
)
from searcher import (
    _job_location_from_json_ld,
    _protect_browser_network,
    _structured_job_metadata,
    _metadata_from_text,
    canonicalize_job_url,
    is_specific_job_url,
    is_public_http_url,
    is_useful_job_details,
    provider_alerts_from_health,
    provider_for_url,
)
from source_diagnostics import (
    MAX_COUNTER_VALUE,
    list_source_diagnostics,
    persist_source_diagnostics,
)
from tailor import (
    RESUME_MODE_GUIDANCE,
    RESUME_SECTION_TEMPLATES,
    TailoringResponse,
    apply_resume_section_template,
    finalize_cover_letter,
    tailor_resume_and_cover_letter,
)
from utils import _headquarters_query, _looks_like_us_address, find_us_headquarters, markdown_to_html


class OperationCancellationTests(unittest.TestCase):
    def test_registered_operation_stops_at_checkpoint_and_cleans_temporary_file(self) -> None:
        operation_id = str(uuid4())
        with tempfile.TemporaryDirectory() as temp_dir:
            temporary_file = Path(temp_dir) / "staged.pdf"
            temporary_file.write_bytes(b"staged")
            token = operations_module.start_operation(operation_id)
            token.track_temporary_file(temporary_file)
            self.assertTrue(operations_module.request_cancellation(operation_id))
            with self.assertRaises(operations_module.OperationCancelled):
                token.checkpoint()
            operations_module.finish_operation(token)
            self.assertFalse(temporary_file.exists())
            self.assertFalse(operations_module.request_cancellation(operation_id))


class AIProviderAbstractionTests(unittest.TestCase):
    def test_openai_responses_request_uses_strict_schema_and_keeps_key_out_of_body(self) -> None:
        output = {
            "output": [{
                "type": "message",
                "content": [{
                    "type": "output_text",
                    "text": json.dumps({"ready": True}),
                }],
            }]
        }

        class FakeResponse:
            def __enter__(self):
                return self

            def __exit__(self, *_args):
                return False

            def read(self, _limit):
                return json.dumps(output).encode("utf-8")

        settings = AIProviderSettings("openai", "gpt-5-mini", "private-test-key")
        with patch.object(ai_providers_module.urllib.request, "urlopen", return_value=FakeResponse()) as open_url:
            result = generate_structured(settings, "Return ready.", CapabilityResponse)

        request = open_url.call_args.args[0]
        payload = json.loads(request.data.decode("utf-8"))
        self.assertEqual({"ready": True}, result)
        self.assertEqual("https://api.openai.com/v1/responses", request.full_url)
        self.assertEqual("Bearer private-test-key", request.get_header("Authorization"))
        self.assertNotIn("private-test-key", request.data.decode("utf-8"))
        self.assertEqual("json_schema", payload["text"]["format"]["type"])
        self.assertTrue(payload["text"]["format"]["strict"])
        self.assertFalse(payload["text"]["format"]["schema"]["additionalProperties"])

    def test_openai_authentication_error_is_provider_specific_and_redacted(self) -> None:
        settings = AIProviderSettings("openai", "gpt-5-mini", "private-test-key")
        response_error = ai_providers_module.urllib.error.HTTPError(
            "https://api.openai.com/v1/responses", 401, "Unauthorized private-test-key", None, None
        )
        with (
            patch.object(ai_providers_module.urllib.request, "urlopen", side_effect=response_error),
            self.assertRaises(AIProviderError) as raised,
        ):
            generate_structured(settings, "Return ready.", CapabilityResponse)

        self.assertEqual("authentication", raised.exception.code)
        self.assertIn("OpenAI rejected", str(raised.exception))
        self.assertNotIn("private-test-key", str(raised.exception))

    def test_openai_ocr_sends_one_bounded_image_and_keeps_key_out_of_body(self) -> None:
        output = {
            "output": [{
                "type": "message",
                "content": [{"type": "output_text", "text": "# Candidate\n\n- Result"}],
            }]
        }

        class FakeResponse:
            def __enter__(self):
                return self

            def __exit__(self, *_args):
                return False

            def read(self, _limit):
                return json.dumps(output).encode("utf-8")

        settings = AIProviderSettings("openai", "gpt-5-mini", "private-ocr-key")
        checkpoints = []
        with patch.object(ai_providers_module.urllib.request, "urlopen", return_value=FakeResponse()) as open_url:
            result = extract_text_from_images(settings, [b"png-test-data"], lambda: checkpoints.append(True))

        request = open_url.call_args.args[0]
        payload = json.loads(request.data.decode("utf-8"))
        image_url = payload["input"][0]["content"][1]["image_url"]
        self.assertEqual(["# Candidate\n\n- Result"], result)
        self.assertTrue(image_url.startswith("data:image/png;base64,"))
        self.assertNotIn("private-ocr-key", request.data.decode("utf-8"))
        self.assertEqual(2, len(checkpoints))

    def test_profile_settings_select_the_openai_key_without_exposing_other_secrets(self) -> None:
        settings = settings_from_profile({
            "ai_provider": "openai",
            "ai_model": "gpt-5-mini",
            "gemini_api_key": "gemini-secret",
            "openai_api_key": "openai-secret",
        })
        self.assertEqual("openai", settings.provider)
        self.assertEqual("gpt-5-mini", settings.model)
        self.assertEqual("openai-secret", settings.api_key)

    def test_keyword_extraction_routes_through_selected_provider(self) -> None:
        with patch.object(
            searcher_module,
            "generate_structured",
            return_value={"titles": ["Data Architect", "Technology Executive"]},
        ) as provider_call:
            titles = searcher_module.extract_search_keywords_from_resume(
                "Evidence-based resume",
                "openai-secret",
                ai_provider="openai",
                ai_model="gpt-5-mini",
            )

        self.assertEqual(["Data Architect", "Technology Executive"], titles)
        settings = provider_call.call_args.args[0]
        self.assertEqual("openai", settings.provider)
        self.assertEqual("gpt-5-mini", settings.model)


class MapsProviderAbstractionTests(unittest.TestCase):
    class FakeResponse:
        def __init__(self, payload):
            self.payload = payload

        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

        def read(self, _limit):
            return json.dumps(self.payload).encode("utf-8")

    def test_google_places_uses_current_post_api_and_keeps_key_out_of_url_and_body(self) -> None:
        response = self.FakeResponse({
            "places": [{
                "formattedAddress": "1600 Amphitheatre Parkway, Mountain View, CA 94043, United States",
                "addressComponents": [{"shortText": "US", "types": ["country"]}],
            }],
        })
        with patch.object(maps_providers_module.urllib.request, "urlopen", return_value=response) as open_url:
            result = lookup_headquarters(
                MapsProviderSettings("google", "private-maps-key"),
                "Example Co",
                prefer_us=True,
            )

        request = open_url.call_args.args[0]
        body = request.data.decode("utf-8")
        headers = {key.lower(): value for key, value in request.header_items()}
        self.assertEqual("POST", request.get_method())
        self.assertEqual(maps_providers_module.GOOGLE_PLACES_SEARCH_URL, request.full_url)
        self.assertEqual("private-maps-key", headers["x-goog-api-key"])
        self.assertNotIn("private-maps-key", request.full_url)
        self.assertNotIn("private-maps-key", body)
        self.assertIn("places.formattedAddress", headers["x-goog-fieldmask"])
        self.assertEqual("US", result.country_code)
        self.assertEqual("Google Maps", result.attribution)

    def test_openstreetmap_identifies_request_filters_us_and_returns_attribution(self) -> None:
        response = self.FakeResponse([{
            "display_name": "Example Co, Dayton, Ohio, United States",
            "address": {"country_code": "us"},
        }])
        maps_providers_module._last_nominatim_request_at = 0.0
        with (
            patch.object(maps_providers_module.config, "get_nominatim_base_url", return_value="https://nominatim.example.test"),
            patch.object(maps_providers_module.urllib.request, "urlopen", return_value=response) as open_url,
        ):
            result = lookup_headquarters(
                MapsProviderSettings("openstreetmap"),
                "Example Co",
                prefer_us=True,
            )

        request = open_url.call_args.args[0]
        self.assertIn("countrycodes=us", request.full_url)
        self.assertIn("format=jsonv2", request.full_url)
        self.assertIn("JobApplierAgent", request.get_header("User-agent"))
        self.assertEqual("openstreetmap", result.source)
        self.assertEqual("US", result.country_code)
        self.assertEqual("© OpenStreetMap contributors", result.attribution)

    def test_maps_provider_readiness_requires_only_google_credentials(self) -> None:
        with patch.object(maps_providers_module.config, "get_google_maps_api_key", return_value=""):
            self.assertFalse(maps_provider_ready({"maps_provider": "google", "google_maps_api_key": ""}))
            self.assertTrue(maps_provider_ready({"maps_provider": "openstreetmap", "google_maps_api_key": ""}))

    def test_new_profile_defaults_select_openstreetmap(self) -> None:
        self.assertEqual("openstreetmap", maps_providers_module.normalize_maps_provider(None))
        profile = app_module.ProfileUpdate(
            name="Candidate",
            email="candidate@example.test",
            phone="",
            github="",
            linkedin="",
            website="",
            base_resume_text="Resume",
        )
        self.assertEqual("openstreetmap", profile.maps_provider)
        html = (Path(__file__).parent / "static" / "index.html").read_text(encoding="utf-8")
        self.assertIn('<option value="openstreetmap" selected>', html)

    def test_headquarters_fallback_uses_selected_ai_provider_and_requires_verification(self) -> None:
        with (
            patch.object(maps_providers_module, "lookup_headquarters", return_value=HeadquartersResult()),
            patch.object(
                maps_providers_module,
                "generate_structured",
                return_value={"address": "1 Main St, Dayton, OH 45402, United States", "country_code": "US", "verified": True},
            ) as generate,
        ):
            result = resolve_headquarters(
                MapsProviderSettings("openstreetmap"),
                AIProviderSettings("openai", "gpt-5-mini", "saved-openai-key"),
                "Example Co",
            )

        self.assertEqual("ai_openai", result.source)
        self.assertIn("verify", result.warning.lower())
        self.assertEqual("openai", generate.call_args.args[0].provider)

    def test_lookup_cache_retains_openstreetmap_but_not_google_places_content(self) -> None:
        connection = sqlite3.connect(":memory:")
        connection.row_factory = sqlite3.Row
        connection.execute("""
            CREATE TABLE headquarters_cache (
                cache_key TEXT PRIMARY KEY, provider TEXT, address TEXT,
                country_code TEXT, attribution TEXT, resolved_at TEXT
            )
        """)
        class NonClosingConnection:
            def execute(self, *args, **kwargs):
                return connection.execute(*args, **kwargs)

            def commit(self):
                connection.commit()

            def close(self):
                pass

        with patch.object(app_module, "get_db_connection", return_value=NonClosingConnection()):
            app_module._cache_headquarters(
                "google-key",
                HeadquartersResult("Google address", "google", "Google Maps", "US"),
                "google",
            )
        self.assertEqual(0, connection.execute("SELECT COUNT(*) FROM headquarters_cache").fetchone()[0])

        with patch.object(app_module, "get_db_connection", return_value=NonClosingConnection()):
            app_module._cache_headquarters(
                "osm-key",
                HeadquartersResult("OSM address", "openstreetmap", "© OpenStreetMap contributors", "US"),
                "openstreetmap",
            )
        self.assertEqual(1, connection.execute("SELECT COUNT(*) FROM headquarters_cache").fetchone()[0])
        connection.close()


class ApplicationMaterialsSafetyTests(unittest.TestCase):
    def test_cover_letter_is_persisted_separately_as_utf8_text(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            cover_path = Path(persist_cover_letter(12, "Dear Hiring Team,\n\nHello.", Path(temp_dir)))
            self.assertEqual("cover_letter_12.txt", cover_path.name)
            self.assertEqual("Dear Hiring Team,\n\nHello.\n", cover_path.read_text(encoding="utf-8"))
            self.assertEqual(cover_path, resolve_output_file(cover_path, ".txt", Path(temp_dir)))

    def test_material_resolution_rejects_files_outside_output_directory(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir, tempfile.TemporaryDirectory() as other_dir:
            outside = Path(other_dir) / "resume.pdf"
            outside.write_bytes(b"pdf")
            with self.assertRaisesRegex(ValueError, "invalid"):
                resolve_output_file(outside, ".pdf", Path(temp_dir))

    def test_download_names_are_safe_and_readable(self) -> None:
        filename = material_download_name("Example / Corp", "VP: Data & AI", "cover-letter", ".txt")
        self.assertEqual("example-corp-vp-data-ai-cover-letter.txt", filename)
        self.assertNotIn("/", filename)

    def test_stopped_pdf_regeneration_preserves_previous_materials(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            database_path = root / "materials.db"
            resume_path = root / "tailored_resume_1.pdf"
            cover_path = root / "cover_letter_1.txt"
            resume_path.write_bytes(b"old pdf")
            cover_path.write_text("Old letter\n", encoding="utf-8")
            connection = sqlite3.connect(database_path)
            connection.executescript("""
                CREATE TABLE profile (id INTEGER PRIMARY KEY, resume_mode TEXT);
                CREATE TABLE applications (
                    id INTEGER PRIMARY KEY, job_id INTEGER, tailored_resume_text TEXT,
                    tailored_resume_path TEXT, cover_letter_path TEXT, cover_letter TEXT
                );
                INSERT INTO profile VALUES (1, 'general_professional');
            """)
            connection.execute(
                "INSERT INTO applications VALUES (1, 1, ?, ?, ?, ?)",
                ("Old resume", str(resume_path), str(cover_path), "Old letter"),
            )
            connection.commit()
            connection.close()

            def connection_factory():
                database = sqlite3.connect(database_path)
                database.row_factory = sqlite3.Row
                return database

            operation_id = str(uuid4())

            def stop_after_render(markdown, output_path, max_pages=2):
                Path(output_path).write_bytes(b"new staged pdf")
                self.assertTrue(app_module.cancel_operation(operation_id)["active"])
                return {"page_count": 1, "compact": False}

            with (
                patch.object(app_module, "get_db_connection", side_effect=connection_factory),
                patch.object(app_module.config, "OUTPUT_DIR", str(root)),
                patch.object(app_module, "apply_resume_section_template", side_effect=lambda text, mode: text),
                patch.object(app_module, "finalize_cover_letter", return_value="New letter"),
                patch.object(app_module, "generate_resume_pdf", side_effect=stop_after_render),
            ):
                result = app_module.update_tailored_details(
                    1,
                    app_module.MaterialsUpdateRequest(
                        tailored_resume="New resume",
                        cover_letter="New letter",
                    ),
                    operation_id,
                )

            self.assertTrue(result["cancelled"])
            self.assertEqual(b"old pdf", resume_path.read_bytes())
            self.assertEqual("Old letter\n", cover_path.read_text(encoding="utf-8"))
            connection = connection_factory()
            stored = connection.execute(
                "SELECT tailored_resume_text, cover_letter FROM applications WHERE job_id = 1"
            ).fetchone()
            connection.close()
            self.assertEqual("Old resume", stored["tailored_resume_text"])
            self.assertEqual("Old letter", stored["cover_letter"])
    def test_cover_letter_date_placeholders_are_resolved(self) -> None:
        expected = "August 3, 2026"
        for placeholder in ("[Date]", "{DATE}", "<date>"):
            with self.subTest(placeholder=placeholder):
                result = finalize_cover_letter(
                    f"Candidate Address\n\n{placeholder}\n\nDear Hiring Team,",
                    date(2026, 8, 3),
                )
                self.assertIn(expected, result)
                self.assertNotIn(placeholder, result)

    def test_downloads_and_manual_open_do_not_mark_job_applied(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            database_path = root / "materials.db"
            resume_path = root / "tailored_resume_1.pdf"
            resume_path.write_bytes(b"%PDF-test")
            cover_path = persist_cover_letter(1, "Dear Hiring Team,", root)

            connection = sqlite3.connect(database_path)
            connection.executescript("""
                CREATE TABLE jobs (
                    id INTEGER PRIMARY KEY, company TEXT, title TEXT, url TEXT,
                    source TEXT, status TEXT
                );
                CREATE TABLE applications (
                    id INTEGER PRIMARY KEY, job_id INTEGER, company TEXT, position TEXT,
                    tailored_resume_path TEXT, cover_letter_path TEXT, cover_letter TEXT,
                    status TEXT
                );
            """)
            connection.execute(
                "INSERT INTO jobs VALUES (1, ?, ?, ?, ?, ?)",
                ("Example Corp", "VP Data", "https://jobs.example.test/1", "lever", "tailored"),
            )
            connection.execute(
                "INSERT INTO applications VALUES (1, 1, ?, ?, ?, ?, ?, ?)",
                ("Example Corp", "VP Data", str(resume_path), cover_path, "Dear Hiring Team,", "tailored"),
            )
            connection.commit()
            connection.close()

            def connection_factory():
                test_connection = sqlite3.connect(database_path)
                test_connection.row_factory = sqlite3.Row
                return test_connection

            with (
                patch.object(app_module, "get_db_connection", connection_factory),
                patch.object(materials_module.config, "OUTPUT_DIR", root),
            ):
                resume_response = app_module.download_tailored_resume(1)
                cover_response = app_module.download_cover_letter(1)
                redirect_response = app_module.open_manual_application(1)

            self.assertIn("example-corp-vp-data-resume.pdf", resume_response.headers["content-disposition"])
            self.assertIn("example-corp-vp-data-cover-letter.txt", cover_response.headers["content-disposition"])
            self.assertEqual("https://jobs.example.test/1", redirect_response.headers["location"])
            connection = connection_factory()
            self.assertEqual("tailored", connection.execute("SELECT status FROM jobs WHERE id = 1").fetchone()[0])
            self.assertEqual("tailored", connection.execute("SELECT status FROM applications WHERE job_id = 1").fetchone()[0])
            connection.close()

class ResumeDocumentInteroperabilityTests(unittest.TestCase):
    def test_docx_import_preserves_semantic_headings_lists_and_table_text(self) -> None:
        document = WordDocument()
        document.add_heading("Candidate Name", level=1)
        document.add_heading("Experience", level=2)
        document.add_paragraph("Delivered measurable results.", style="List Bullet")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Email"
        table.cell(0, 1).text = "candidate@example.com"
        buffer = io.BytesIO()
        document.save(buffer)

        imported = import_resume_document("resume.docx", buffer.getvalue())

        self.assertEqual("docx", imported.source_format)
        self.assertIn("# Candidate Name", imported.text)
        self.assertIn("## Experience", imported.text)
        self.assertIn("- Delivered measurable results.", imported.text)
        self.assertIn("Email | candidate@example.com", imported.text)
        self.assertFalse(imported.ocr_used)

    def test_text_pdf_import_remains_local(self) -> None:
        document = pymupdf.open()
        page = document.new_page()
        page.insert_text(
            (72, 72),
            "Candidate Name - Experience delivering measurable business outcomes across several roles.",
        )
        content = document.tobytes()
        document.close()

        imported = import_resume_document("resume.pdf", content)

        self.assertEqual("pdf", imported.source_format)
        self.assertEqual(1, imported.page_count)
        self.assertFalse(imported.ocr_used)
        self.assertIn("Candidate Name", imported.text)

    def test_scanned_pdf_requires_explicit_ocr_consent(self) -> None:
        document = pymupdf.open()
        document.new_page()
        content = document.tobytes()
        document.close()

        with self.assertRaisesRegex(ResumeDocumentError, "Select the AI OCR option"):
            import_resume_document("resume.pdf", content)

        received = []

        def ocr(images, checkpoint):
            received.extend(images)
            if checkpoint:
                checkpoint()
            return ["# Candidate Name\n\n## Experience\n\n- OCR result"]

        imported = import_resume_document(
            "resume.pdf",
            content,
            allow_ocr=True,
            ocr_images=ocr,
            checkpoint=lambda: None,
        )
        self.assertTrue(imported.ocr_used)
        self.assertEqual(1, len(received))
        self.assertTrue(received[0].startswith(b"\x89PNG"))
        self.assertIn("OCR result", imported.text)

    def test_import_bounds_text_and_large_page_ocr_renders(self) -> None:
        with self.assertRaisesRegex(ResumeDocumentError, "200,000 characters"):
            import_resume_document("resume.txt", b"a" * (MAX_IMPORTED_RESUME_CHARACTERS + 1))

        document = pymupdf.open()
        document.new_page(width=10_000, height=10_000)
        content = document.tobytes()
        document.close()
        received = []

        def ocr(images, _checkpoint):
            received.extend(images)
            return ["Bounded OCR result with enough readable resume text for the editor."]

        import_resume_document("resume.pdf", content, allow_ocr=True, ocr_images=ocr)
        rendered = pymupdf.Pixmap(received[0])
        self.assertLessEqual(max(rendered.width, rendered.height), MAX_OCR_RENDER_DIMENSION)
        self.assertLessEqual(rendered.width * rendered.height, MAX_OCR_RENDER_PIXELS)

    def test_accessible_docx_uses_real_styles_language_and_hyperlink_relationships(self) -> None:
        content = build_accessible_resume_docx(
            "# Candidate Name\n\ncandidate@example.com | (937) 555-0123 | https://example.com\n\n"
            "## Experience\n\n- Delivered **measurable** results.",
            candidate_name="Candidate Name",
            title="Candidate Resume",
        )
        loaded = WordDocument(io.BytesIO(content))
        styles = [paragraph.style.name for paragraph in loaded.paragraphs]
        relationships = [
            relationship.target_ref
            for relationship in loaded.part.rels.values()
            if relationship.reltype.endswith("/hyperlink")
        ]

        self.assertEqual("Candidate Resume", loaded.core_properties.title)
        self.assertEqual("Candidate Name", loaded.core_properties.author)
        self.assertIn("Heading 1", styles)
        self.assertIn("Heading 2", styles)
        self.assertIn("List Bullet", styles)
        self.assertIn("mailto:candidate@example.com", relationships)
        self.assertIn("tel:9375550123", relationships)
        self.assertIn("https://example.com", relationships)
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            styles_xml = archive.read("word/styles.xml").decode("utf-8")
        self.assertIn('w:val="en-US"', styles_xml)

    def test_docx_download_is_generated_from_reviewed_source_with_safe_name(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "materials.db"
            connection = sqlite3.connect(database_path)
            connection.executescript("""
                CREATE TABLE profile (id INTEGER PRIMARY KEY, name TEXT);
                INSERT INTO profile VALUES (1, 'Candidate Name');
                CREATE TABLE jobs (
                    id INTEGER PRIMARY KEY, company TEXT, title TEXT, url TEXT,
                    source TEXT, status TEXT
                );
                INSERT INTO jobs VALUES (
                    1, 'Example Corp', 'Data Leader', 'https://example.test/job', 'manual', 'tailored'
                );
                CREATE TABLE applications (job_id INTEGER PRIMARY KEY, tailored_resume_text TEXT);
                INSERT INTO applications VALUES (
                    1, '# Candidate Name\n\n## Experience\n\n- Evidence-based result.'
                );
            """)
            connection.close()

            def connection_factory():
                test_connection = sqlite3.connect(database_path)
                test_connection.row_factory = sqlite3.Row
                return test_connection

            with patch.object(app_module, "get_db_connection", side_effect=connection_factory):
                response = app_module.download_tailored_resume_docx(1)

        self.assertEqual(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            response.media_type,
        )
        self.assertIn("example-corp-data-leader-resume-accessible.docx", response.headers["content-disposition"])
        loaded = WordDocument(io.BytesIO(response.body))
        self.assertEqual("Candidate Name", loaded.paragraphs[0].text)


class ResumeRenderingTests(unittest.TestCase):
    def test_tailoring_uses_validated_structured_response(self) -> None:
        generated = {
            "tailored_resume": "# Candidate\n\n## Professional Summary\n\nEvidence-based summary.",
            "cover_letter": "August 4, 2026\n\nDear Hiring Team,\n\nI am interested in this role.",
        }
        with patch("tailor.generate_structured", return_value=generated) as provider_call:
            result = tailor_resume_and_cover_letter(
                "# Candidate\n\n## Experience\n\nEvidence.",
                "Director",
                "Example Company",
                "Lead an evidence-based program.",
            )

        self.assertTrue(result["success"])
        self.assertIn("## Professional Summary", result["tailored_resume"])
        self.assertIs(provider_call.call_args.args[2], TailoringResponse)

    def test_tailoring_prompt_includes_only_supplied_professional_evidence(self) -> None:
        generated = {
            "tailored_resume": "# Candidate\n\n## Summary\n\nEvidence-based summary.",
            "cover_letter": "August 8, 2026\n\nDear Hiring Team,\n\nI am interested.",
        }
        evidence = professional_evidence_markdown({
            "skills": "Python\nCloud architecture",
            "portfolio": "Design portfolio — https://example.test/portfolio",
        })
        with patch("tailor.generate_structured", return_value=generated) as provider_call:
            result = tailor_resume_and_cover_letter(
                "# Candidate\n\n## Experience\n\nEvidence.",
                "Architect",
                "Example Company",
                "Lead cloud architecture.",
                resume_mode="it",
                professional_evidence=evidence,
            )

        prompt = provider_call.call_args.args[1]
        self.assertTrue(result["success"])
        self.assertIn("### Skills", prompt)
        self.assertIn("https://example.test/portfolio", prompt)
        self.assertIn("never infer or invent missing dates", prompt)
        self.assertNotIn("### Licenses", prompt)

    def test_markdown_is_escaped_and_supported_tokens_are_rendered(self) -> None:
        rendered = markdown_to_html(
            "# Candidate\n\n---\n\n## Experience\n\n### Role\n\n"
            "- Delivered **measurable** results with *care*.\n\n<script>alert(1)</script>"
        )
        self.assertIn("<hr>", rendered)
        self.assertIn("<strong>measurable</strong>", rendered)
        self.assertIn("<em>care</em>", rendered)
        self.assertNotIn("<script>", rendered)
        self.assertIn("&lt;script&gt;", rendered)
        self.assertIn('<section class="resume-section">', rendered)


class JobFilterFacetTests(unittest.TestCase):
    def test_technical_role_facets_cover_pay_travel_sponsorship_and_clearance(self) -> None:
        facets = derive_job_filter_facets({
            "title": "Cloud Security Director",
            "description": (
                "This remote role participates in an on-call rotation and requires up to 25% travel. "
                "Candidates must hold an active TS/SCI clearance. No visa sponsorship is available."
            ),
            "compensation": "$180k - $220k per year",
            "work_arrangement": "remote",
            "employment_type": "full_time",
        })

        self.assertEqual("annual", facets["compensation_period"])
        self.assertEqual(180_000, facets["compensation_min"])
        self.assertEqual(220_000, facets["compensation_max"])
        self.assertEqual("remote", facets["commute_requirement"])
        self.assertEqual("full_time", facets["employment_type"])
        self.assertEqual(25, facets["travel_percent"])
        self.assertIn("on_call", facets["shift_tags"])
        self.assertEqual("unavailable", facets["sponsorship"])
        self.assertEqual(4, facets["clearance_rank"])

    def test_healthcare_role_facets_cover_hourly_shift_license_and_conditions(self) -> None:
        facets = derive_job_filter_facets({
            "title": "Emergency Department RN",
            "description": (
                "Night shift. Active RN license required. Must lift up to 50 pounds and "
                "stand for extended periods."
            ),
            "compensation": "$42-$55/hr",
            "work_arrangement": "on_site",
            "employment_type": "part_time",
        })

        self.assertEqual("hourly", facets["compensation_period"])
        self.assertEqual(42, facets["compensation_min"])
        self.assertEqual(55, facets["compensation_max"])
        self.assertIn("night", facets["shift_tags"])
        self.assertIn("registered_nurse", facets["license_tags"])
        self.assertTrue(facets["license_required"])
        self.assertTrue(facets["physical_conditions"])
        self.assertIn("lifting", facets["condition_tags"])
        self.assertIn("standing", facets["condition_tags"])

    def test_non_it_trade_role_facets_cover_contract_rate_and_work_conditions(self) -> None:
        facets = derive_job_filter_facets({
            "title": "Commercial Driver",
            "description": (
                "A CDL-A is required. Regular driving and outdoor work. Zero travel outside "
                "the assigned local route. Visa sponsorship is available."
            ),
            "compensation": "$38 per hour",
            "work_arrangement": "on_site",
            "employment_type": "contract",
        })

        self.assertEqual("hourly", facets["compensation_period"])
        self.assertFalse(facets["travel_required"])
        self.assertEqual(0, facets["travel_percent"])
        self.assertEqual("available", facets["sponsorship"])
        self.assertIn("commercial_driver", facets["license_tags"])
        self.assertTrue(facets["physical_conditions"])

    def test_unknown_and_negative_language_do_not_invent_requirements(self) -> None:
        facets = derive_job_filter_facets({
            "title": "Software Account Executive",
            "description": "Sell licensed software with overnight data refreshes and a night mode. No Secret clearance is required.",
            "compensation": "Competitive",
            "work_arrangement": "",
            "employment_type": "",
        })

        self.assertEqual("unknown", facets["compensation_period"])
        self.assertEqual("unknown", facets["commute_requirement"])
        self.assertEqual("unknown", facets["employment_type"])
        self.assertEqual(0, facets["clearance_rank"])
        self.assertFalse(facets["license_required"])
        self.assertEqual([], facets["shift_tags"])
        self.assertIsNone(facets["travel_required"])


class FrontendStartupTests(unittest.TestCase):
    def test_required_startup_controls_exist_in_dashboard_html(self) -> None:
        static_dir = Path(__file__).parent / "static"
        html_source = (static_dir / "index.html").read_text(encoding="utf-8")
        for element_id in (
            "profile-form", "resume-file-upload", "search-form", "refresh-jobs-btn",
            "cleanup-jobs-btn", "refresh-logs-btn", "open-manual-application-btn",
            "archive-untouched-btn", "delete-untouched-btn", "restore-archived-btn",
            "lifecycle-form", "undo-lifecycle-btn", "save-materials-btn",
            "download-resume-btn", "download-resume-docx-btn", "download-cover-letter-btn",
            "saved-search-select", "p-resume-mode",
            "p-ai-provider", "p-ai-model", "p-openai-apikey",
            "p-openai-key-status", "p-openai-key-help", "test-ai-provider-btn",
            "p-maps-provider", "test-maps-provider-btn", "maps-provider-test-status",
            "google-maps-key-group", "openstreetmap-policy",
            "p-prefer-us-headquarters",
            "p-gemini-key-status", "p-gemini-key-help",
            "p-google-key-status", "p-google-key-help",
            "startup-activity-title", "startup-activity-description",
            "lifecycle-applied-calendar-btn", "saved-search-frequency", "provider-alerts",
            "open-job-import-btn", "job-import-modal", "job-import-form",
            "job-import-url", "preview-job-import-btn", "job-import-fields",
            "job-import-company", "job-import-title", "job-import-description",
            "save-job-import-btn", "suppression-count", "suppression-list",
            "suppression-empty", "clear-all-suppressions-btn",
            "open-source-diagnostics-btn", "source-diagnostics-count",
            "source-diagnostics-modal", "source-diagnostics-list",
            "source-diagnostics-empty", "clear-source-diagnostics-btn",
            "export-source-diagnostics-btn",
            "loading-actions", "stop-loading-btn",
            "base-resume-select", "p-resume-name", "new-base-resume-btn",
            "duplicate-base-resume-btn", "base-resume-history-btn",
            "delete-base-resume-btn", "resume-ocr-consent", "base-resume-history-modal",
            "base-resume-version-list", "restore-base-resume-version-btn",
            "professional-evidence-editor", "professional-evidence-count",
            "evidence-mode-guidance", "p-evidence-skills", "p-evidence-projects",
            "p-evidence-portfolio", "p-evidence-licenses",
            "p-evidence-certifications", "p-evidence-work-samples",
            "advanced-job-filters", "job-active-filter-count",
            "job-employment-filter", "job-commute-filter",
            "job-min-annual-compensation", "job-min-hourly-rate",
            "job-shift-filter", "job-max-travel-filter",
            "job-sponsorship-filter", "job-clearance-filter",
            "job-license-filter", "job-conditions-filter",
            "job-include-unknown", "reset-advanced-job-filters",
            "application-insights-card", "application-insights-dimension",
            "application-insights-summary", "application-insights-table",
            "application-insights-body", "application-insights-note",
        ):
            with self.subTest(element_id=element_id):
                self.assertIn(f'id="{element_id}"', html_source)
        self.assertIn("Checking AI Provider", html_source)
        self.assertIn("app.js?v=20260808-13", html_source)
        self.assertIn("index.css?v=20260808-13", html_source)

    def test_launchers_require_the_current_backend_build(self) -> None:
        project_dir = Path(__file__).parent.parent
        for launcher in ("run.bat", "run.ps1"):
            with self.subTest(launcher=launcher):
                source = (project_dir / launcher).read_text(encoding="utf-8")
                self.assertIn("/api/version", source)
                self.assertIn("20260808.13", source)

    def test_advanced_job_filters_are_local_and_do_not_change_match_scores(self) -> None:
        project_dir = Path(__file__).parent.parent
        app_source = (project_dir / "backend" / "app.py").read_text(encoding="utf-8")
        script_source = (project_dir / "backend" / "static" / "app.js").read_text(encoding="utf-8")
        self.assertIn('job["filter_facets"] = derive_job_filter_facets(job)', app_source)
        self.assertIn("function jobMatchesAdvancedFilters", script_source)
        self.assertIn("jobs = jobs.filter(job => jobMatchesAdvancedFilters", script_source)
        self.assertNotIn("match_score =", script_source[script_source.index("function jobMatchesAdvancedFilters"):script_source.index("function activeAdvancedJobFilterCount")])

    def test_manual_application_flow_replaces_browser_submission(self) -> None:
        project_dir = Path(__file__).parent.parent
        app_source = (project_dir / "backend" / "app.py").read_text(encoding="utf-8")
        html_source = (project_dir / "backend" / "static" / "index.html").read_text(encoding="utf-8")
        script_source = (project_dir / "backend" / "static" / "app.js").read_text(encoding="utf-8")
        self.assertFalse((project_dir / "backend" / "applier.py").exists())
        self.assertNotIn('@app.post("/api/jobs/{job_id}/apply")', app_source)
        self.assertIn('@app.get("/api/jobs/{job_id}/apply-manually")', app_source)
        self.assertNotIn("Submit Application Now", html_source)
        self.assertNotIn("triggerApplicationSubmission", script_source)
        self.assertIn("Apply Manually", script_source)

    def test_base_resume_controls_are_bound_to_versioned_workflows(self) -> None:
        app_source = (Path(__file__).parent / "app.py").read_text(encoding="utf-8")
        script_source = (Path(__file__).parent / "static" / "app.js").read_text(encoding="utf-8")
        self.assertIn('bindEvent(newBaseResumeBtn, "click", startNewBaseResume)', script_source)
        self.assertIn('bindEvent(duplicateBaseResumeBtn, "click", beginBaseResumeCopy)', script_source)
        self.assertIn('bindEvent(baseResumeHistoryBtn, "click", showBaseResumeHistory)', script_source)
        self.assertIn('bindEvent(deleteBaseResumeBtn, "click", removeBaseResume)', script_source)
        self.assertIn("professionalEvidencePreview(version.professional_evidence)", script_source)
        self.assertIn("ROLE-SPECIFIC PROFESSIONAL EVIDENCE", script_source)
        self.assertIn("professional_evidence: readProfessionalEvidence()", script_source)
        self.assertIn("profile.professional_evidence.model_dump()", app_source)

    def test_document_import_and_docx_download_controls_are_bound(self) -> None:
        script_source = (Path(__file__).parent / "static" / "app.js").read_text(encoding="utf-8")
        self.assertIn('formData.append("allow_ocr",', script_source)
        self.assertIn("/materials/resume.docx", script_source)
        self.assertIn('bindEvent(resumeFileUpload, "change", handleResumeUpload)', script_source)

    def test_launchers_use_the_configurable_default_port(self) -> None:
        project_dir = Path(__file__).parent.parent
        powershell_source = (project_dir / "run.ps1").read_text(encoding="utf-8")
        batch_source = (project_dir / "run.bat").read_text(encoding="utf-8")

        self.assertIn("[int]$Port = 8001", powershell_source)
        self.assertIn("--port $Port", powershell_source)
        self.assertIn('set "JOBAPPLIER_REQUESTED_PORT=8001"', batch_source)
        self.assertIn('set "PORT=%JOBAPPLIER_REQUESTED_PORT%"', batch_source)
        self.assertIn("--port %PORT%", batch_source)
        self.assertNotIn("8000", powershell_source)
        self.assertNotIn("8000", batch_source)

    def test_job_applier_brand_assets_and_manifest_are_wired(self) -> None:
        static_dir = Path(__file__).parent / "static"
        html_source = (static_dir / "index.html").read_text(encoding="utf-8")
        manifest_source = (static_dir / "site.webmanifest").read_text(encoding="utf-8")
        for asset in (
            "icons/favicon.ico", "icons/favicon.svg", "icons/apple-touch-icon.png",
            "icons/safari-pinned-tab.svg", "icons/icon-192.png", "icons/icon-512.png",
        ):
            with self.subTest(asset=asset):
                self.assertTrue((static_dir / asset).is_file())
        self.assertIn('rel="manifest" href="/static/site.webmanifest"', html_source)
        self.assertIn('src="/static/icons/favicon.svg"', html_source)
        self.assertIn('"name": "Job Applier Agent"', manifest_source)

    def test_cleanup_footer_has_scoped_responsive_actions(self) -> None:
        static_dir = Path(__file__).parent / "static"
        html_source = (static_dir / "index.html").read_text(encoding="utf-8")
        css_source = (static_dir / "index.css").read_text(encoding="utf-8")
        self.assertIn('class="modal-footer cleanup-footer"', html_source)
        self.assertIn('class="cleanup-count-badge"', html_source)
        self.assertIn('btn btn-primary cleanup-action-btn" id="archive-untouched-btn"', html_source)
        self.assertIn(".cleanup-footer .modal-actions", css_source)
        self.assertIn("@media (max-width: 620px)", css_source)

    def test_cleanup_actions_have_click_through_guard(self) -> None:
        script_source = (Path(__file__).parent / "static" / "app.js").read_text(encoding="utf-8")
        self.assertIn("cleanupActionsReady", script_source)
        self.assertIn("}, 500);", script_source)

    def test_job_and_log_rows_do_not_interpolate_untrusted_html(self) -> None:
        script_source = (Path(__file__).parent / "static" / "app.js").read_text(encoding="utf-8")
        self.assertNotIn("tr.innerHTML", script_source)
        self.assertIn("safeHttpUrl", script_source)

    def test_all_resume_modes_have_prompt_guidance(self) -> None:
        expected = {"it", "technical_executive", "general_professional", "federal", "healthcare", "education", "sales", "trades_operations", "academic_cv", "cover_letter"}
        self.assertEqual(expected, set(RESUME_MODE_GUIDANCE))
        self.assertEqual(expected, set(RESUME_SECTION_TEMPLATES))

    def test_dual_mode_date_entry_is_validated(self) -> None:
        script_source = (Path(__file__).parent / "static" / "app.js").read_text(encoding="utf-8")
        self.assertIn("parseDateEntry", script_source)
        self.assertIn("showPicker", script_source)

    def test_profile_api_source_does_not_return_plaintext_secrets_or_enable_cors(self) -> None:
        app_source = (Path(__file__).parent / "app.py").read_text(encoding="utf-8")
        self.assertNotIn("CORSMiddleware", app_source)
        self.assertIn('result.pop("gemini_api_key"', app_source)
        self.assertIn('result.pop("openai_api_key"', app_source)
        self.assertIn('result.pop("google_maps_api_key"', app_source)

    def test_profile_ui_exposes_privacy_safe_key_status_and_replacement_guidance(self) -> None:
        static_dir = Path(__file__).parent / "static"
        html_source = (static_dir / "index.html").read_text(encoding="utf-8")
        script_source = (static_dir / "app.js").read_text(encoding="utf-8")
        self.assertIn('role="status" aria-live="polite"', html_source)
        self.assertIn("updateSecretStatuses", script_source)
        self.assertIn('status.textContent = configured ? "Saved" : "Not saved"', script_source)
        self.assertIn("Leave this field blank to keep the current key", script_source)
        self.assertNotRegex(script_source, r"profile\.google_maps_api_key(?!_configured)")
        self.assertNotRegex(script_source, r"profile\.gemini_api_key(?!_configured)")
        self.assertNotRegex(script_source, r"profile\.openai_api_key(?!_configured)")

    def test_manual_import_ui_preserves_unscored_jobs(self) -> None:
        script_source = (Path(__file__).parent / "static" / "app.js").read_text(encoding="utf-8")
        self.assertIn("/api/jobs/import/preview", script_source)
        self.assertIn("/api/jobs/import", script_source)
        self.assertIn('job.match_score === null', script_source)
        self.assertIn('"Unscored"', script_source)
        self.assertIn("showCancellableLoading", script_source)
        self.assertIn("revealJobId", script_source)
        self.assertIn("job-row-revealed", script_source)

    def test_expected_search_limits_render_as_informational_notes(self) -> None:
        script_source = (Path(__file__).parent / "static" / "app.js").read_text(encoding="utf-8")
        self.assertIn('["stale_postings", "partial_results"]', script_source)
        self.assertIn('needsAttention ? "Some job sources may need attention" : "Search notes"', script_source)

    def test_source_notices_are_dismissible_with_persistent_history_access(self) -> None:
        static_dir = Path(__file__).parent / "static"
        html_source = (static_dir / "index.html").read_text(encoding="utf-8")
        script_source = (static_dir / "app.js").read_text(encoding="utf-8")
        css_source = (static_dir / "index.css").read_text(encoding="utf-8")
        self.assertIn('id="provider-alerts" role="status" aria-live="polite"', html_source)
        self.assertIn('role="dialog" aria-modal="true" aria-labelledby="source-diagnostics-title"', html_source)
        self.assertIn("dismissProviderAlerts", script_source)
        self.assertIn("openSourceDiagnosticsBtn.focus()", script_source)
        self.assertIn("showSourceDiagnostics", script_source)
        self.assertIn("loadSourceDiagnostics", script_source)
        self.assertIn("createElement(\"details\"", script_source)
        self.assertNotIn("sourceDiagnosticsList.innerHTML", script_source)
        self.assertIn(".provider-alert-dismiss:focus-visible", css_source)
        self.assertIn("@media (max-width: 620px)", css_source)

    def test_startup_activity_reflects_loaded_profile_state(self) -> None:
        static_dir = Path(__file__).parent / "static"
        html_source = (static_dir / "index.html").read_text(encoding="utf-8")
        script_source = (static_dir / "app.js").read_text(encoding="utf-8")
        self.assertNotIn("Fill in your profile details to start.", html_source)
        self.assertIn("updateStartupActivity(profile)", script_source)
        self.assertIn('title.textContent = "Profile Loaded"', script_source)
        self.assertIn('title.textContent = "Profile Setup Needed"', script_source)


class DependencyLockTests(unittest.TestCase):
    @staticmethod
    def project_dir() -> Path:
        return Path(__file__).parent.parent

    def dependency_files(self) -> tuple[Path, Path]:
        backend_dir = self.project_dir() / "backend"
        policy = Path(os.environ.get("JOBAPPLIER_TEST_REQUIREMENTS_IN", backend_dir / "requirements.in"))
        lock = Path(os.environ.get("JOBAPPLIER_TEST_REQUIREMENTS_LOCK", backend_dir / "requirements.txt"))
        return policy, lock

    def test_direct_dependency_policy_is_fully_hash_locked(self) -> None:
        policy, lock = self.dependency_files()
        direct_source = policy.read_text(encoding="utf-8")
        lock_source = lock.read_text(encoding="utf-8")

        direct_names = {
            re.split(r"[<>=!~\[]", line, maxsplit=1)[0].strip().lower().replace("_", "-")
            for line in direct_source.splitlines()
            if line.strip() and not line.lstrip().startswith("#")
        }
        locked_blocks: dict[str, list[str]] = {}
        current_name = None
        for line in lock_source.splitlines():
            match = re.match(r"^([A-Za-z0-9_.-]+)==([^\s\\]+)\s*\\?$", line)
            if match:
                current_name = match.group(1).lower().replace("_", "-")
                locked_blocks[current_name] = [line]
            elif current_name:
                locked_blocks[current_name].append(line)

        self.assertEqual(
            {
                "exceptiongroup", "fastapi", "google-genai", "playwright", "posthog",
                "pydantic", "pymupdf", "python-docx", "python-multipart", "uvicorn",
            },
            direct_names,
        )
        self.assertTrue(direct_names.issubset(locked_blocks))
        self.assertNotIn("jinja2", locked_blocks)
        self.assertNotIn("duckduckgo-search", locked_blocks)
        self.assertEqual("1.3.1", locked_requirements(lock)["exceptiongroup"])
        for package, block in locked_blocks.items():
            with self.subTest(package=package):
                self.assertRegex(block[0], rf"^{re.escape(package)}==", msg=block[0].lower())
                self.assertIn("--hash=sha256:", "\n".join(block))

    def test_lock_rejects_mutable_or_remote_requirements(self) -> None:
        _, application_lock = self.dependency_files()
        tool_lock = self.project_dir() / "scripts" / "dependency-tools.txt"
        for lock in (application_lock, tool_lock):
            lock_source = lock.read_text(encoding="utf-8")
            install_lines = [line.strip() for line in lock_source.splitlines() if line and not line[0].isspace() and not line.startswith("#")]
            self.assertTrue(install_lines)
            for line in install_lines:
                with self.subTest(lock=lock.name, line=line):
                    self.assertRegex(line, r"^[A-Za-z0-9_.-]+==[^\s\\]+\s*\\?$")
                    self.assertNotRegex(line.lower(), r"(?:https?://|git\+|file:|\s@\s|^-e\s|\.\.[/\\])")

    def test_dependency_compiler_is_fully_hash_locked(self) -> None:
        tool_input = (self.project_dir() / "scripts" / "dependency-tools.in").read_text(encoding="utf-8")
        tool_lock = self.project_dir() / "scripts" / "dependency-tools.txt"
        self.assertEqual(
            {"pip", "pip-tools", "tomli"},
            {
                re.split(r"[<>=!~\[]", line, maxsplit=1)[0].strip().lower()
                for line in tool_input.splitlines()
                if line.strip() and not line.lstrip().startswith("#")
            },
        )
        locked = locked_requirements(tool_lock)
        self.assertEqual("26.1.2", locked["pip"])
        self.assertEqual("7.6.0", locked["pip-tools"])
        self.assertEqual("2.4.1", locked["tomli"])
        tool_source = tool_lock.read_text(encoding="utf-8")
        blocks: dict[str, list[str]] = {}
        current_name = None
        for line in tool_source.splitlines():
            match = re.match(r"^([A-Za-z0-9_.-]+)==([^\s\\]+)\s*\\?$", line)
            if match:
                current_name = match.group(1).lower().replace("_", "-")
                blocks[current_name] = [line]
            elif current_name:
                blocks[current_name].append(line)
        for package in locked:
            self.assertIn("--hash=sha256:", "\n".join(blocks[package]))

    def test_launchers_reconcile_and_stamp_the_reviewed_lock(self) -> None:
        for launcher in ("run.bat", "run.ps1"):
            with self.subTest(launcher=launcher):
                source = (self.project_dir() / launcher).read_text(encoding="utf-8")
                lower_source = source.lower()
                self.assertIn("requirements.in", lower_source)
                self.assertIn("requirements.txt", lower_source)
                self.assertIn("dependency_lock.py", lower_source)
                self.assertIn("--lock", lower_source)
                self.assertIn("--require-hashes", lower_source)
                self.assertIn("pip check", lower_source)
                self.assertIn("playwright install chromium", lower_source)
                self.assertNotIn("pip install --upgrade pip", lower_source)
                self.assertLess(
                    lower_source.index("playwright install chromium"),
                    lower_source.index("write --stamp"),
                )
                self.assertIn(".venv-previous-", lower_source)
                self.assertNotIn("venv --clear", lower_source)
                self.assertLess(lower_source.index(".venv-previous-"), lower_source.index("-m venv"))
                self.assertLess(lower_source.index("--require-hashes"), lower_source.index("write --stamp"))
                self.assertLess(lower_source.index("write --stamp"), lower_source.index("environment repair complete"))
                self.assertIn("stop any running job applier copy with ctrl+c", lower_source)
        powershell_source = (self.project_dir() / "run.ps1").read_text(encoding="utf-8")
        batch_source = (self.project_dir() / "run.bat").read_text(encoding="utf-8")
        self.assertIn("(3, 10, 2)", powershell_source)
        self.assertIn("(3, 10, 2)", batch_source)
        self.assertIn("Start-Job", powershell_source)
        self.assertNotIn("Start-ThreadJob", powershell_source)
        self.assertIn("[int]::TryParse", batch_source)
        self.assertIn("SERVER_EXIT", batch_source)

    def test_environment_manifest_rejects_unlocked_or_wrong_packages(self) -> None:
        _, lock = self.dependency_files()

        class FakeDistribution:
            def __init__(self, name: str, version: str):
                self.metadata = {"Name": name}
                self.version = version

        expected = locked_requirements(lock)
        exact = [FakeDistribution(name, version) for name, version in expected.items()]
        self.assertTrue(environment_matches_lock(lock, exact + [FakeDistribution("pip", "1.0")]))
        self.assertFalse(environment_matches_lock(lock, exact + [FakeDistribution("unreviewed", "1.0")]))
        wrong = exact.copy()
        wrong[0] = FakeDistribution(wrong[0].metadata["Name"], "0.0")
        self.assertFalse(environment_matches_lock(lock, wrong))

    def test_dependency_stamp_detects_policy_and_lock_changes(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            policy = root / "requirements.in"
            lock = root / "requirements.txt"
            stamp = root / "venv" / ".jobapplier-requirements.sha256"
            policy.write_text("fastapi~=0.137.0\n", encoding="utf-8")
            lock.write_text("fastapi==0.137.2 \\\n    --hash=sha256:" + "a" * 64 + "\n", encoding="utf-8")

            self.assertFalse(stamp_is_current(stamp, [policy, lock]))
            first_fingerprint = dependency_fingerprint([policy, lock])
            write_stamp(stamp, [policy, lock])
            self.assertTrue(stamp_is_current(stamp, [policy, lock]))
            self.assertEqual(first_fingerprint, stamp.read_text(encoding="ascii").strip())

            lock.write_text(lock.read_text(encoding="utf-8").replace("0.137.2", "0.137.3"), encoding="utf-8")
            self.assertFalse(stamp_is_current(stamp, [policy, lock]))

    def test_update_workflow_uses_clean_pinned_tooling_and_validation(self) -> None:
        script = (self.project_dir() / "scripts" / "update_dependencies.ps1").read_text(encoding="utf-8")
        self.assertIn("dependency-tools.txt", script)
        self.assertIn("--only-binary=:all:", script)
        self.assertIn("--generate-hashes", script)
        self.assertIn("--require-hashes", script)
        self.assertIn("test_p0_regressions.py", script)
        self.assertIn("test_analytics.py", script)
        self.assertIn("pip check", script)
        self.assertIn("[IO.File]::Replace", script)
        self.assertIn("Find-UpdatePython", script)

    def test_clean_lock_ci_covers_supported_python_versions(self) -> None:
        workflow = (self.project_dir() / ".github" / "workflows" / "dependency-tests.yml").read_text(encoding="utf-8")
        self.assertIn('python-version: ["3.10", "3.12"]', workflow)
        self.assertIn("dependency-tools.txt", workflow)
        self.assertIn(".dependency-tools-venv", workflow)
        self.assertIn("--require-hashes", workflow)
        self.assertIn("test_p0_regressions.py", workflow)
        self.assertIn("test_analytics.py", workflow)


class ProfileSecretPresenceTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.db_path = Path(self.temp_dir.name) / "profile-secrets.db"
        connection = sqlite3.connect(self.db_path)
        connection.executescript("""
            CREATE TABLE profile (
                id INTEGER PRIMARY KEY,
                name TEXT,
                gemini_api_key TEXT,
                openai_api_key TEXT,
                google_maps_api_key TEXT,
                ai_provider TEXT,
                ai_model TEXT,
                maps_provider TEXT
            );
            INSERT INTO profile VALUES (
                1, 'Test Candidate', 'gemini-secret', '', '', 'gemini', 'gemini-2.5-flash', 'google'
            );
        """)
        connection.commit()
        connection.close()

    def tearDown(self) -> None:
        self.temp_dir.cleanup()

    def connection_factory(self):
        connection = sqlite3.connect(self.db_path)
        connection.row_factory = sqlite3.Row
        return connection

    def test_profile_returns_presence_flags_without_secret_values(self) -> None:
        with patch.object(app_module, "get_db_connection", side_effect=self.connection_factory):
            profile = app_module.get_profile()

        self.assertTrue(profile["gemini_api_key_configured"])
        self.assertFalse(profile["openai_api_key_configured"])
        self.assertFalse(profile["google_maps_api_key_configured"])
        self.assertNotIn("gemini_api_key", profile)
        self.assertNotIn("openai_api_key", profile)
        self.assertNotIn("google_maps_api_key", profile)

    def test_openai_secret_update_returns_only_presence_flags(self) -> None:
        with (
            patch.object(app_module, "get_db_connection", side_effect=self.connection_factory),
            patch.object(app_module.config, "OPENAI_API_KEY", ""),
            patch.dict(os.environ, {}, clear=False),
        ):
            os.environ.pop("OPENAI_API_KEY", None)
            result = app_module.update_profile_secrets(
                app_module.ProfileSecretsUpdate(openai_api_key="openai-private")
            )

        connection = self.connection_factory()
        stored = connection.execute("SELECT openai_api_key FROM profile WHERE id = 1").fetchone()[0]
        connection.close()
        self.assertEqual("openai-private", stored)
        self.assertTrue(result["openai_api_key_configured"])
        self.assertNotIn("openai_api_key", result)

    def test_provider_capability_endpoint_uses_saved_provider_and_model(self) -> None:
        connection = self.connection_factory()
        connection.execute(
            "UPDATE profile SET ai_provider = 'openai', ai_model = 'gpt-5-mini', openai_api_key = 'saved-key' WHERE id = 1"
        )
        connection.commit()
        connection.close()
        provider_result = {
            "success": True,
            "provider": "openai",
            "provider_label": "OpenAI",
            "model": "gpt-5-mini",
            "message": "OpenAI is ready for matching and tailoring.",
        }
        with (
            patch.object(app_module, "get_db_connection", side_effect=self.connection_factory),
            patch.object(app_module, "validate_provider_capability", return_value=provider_result) as validate,
        ):
            result = app_module.validate_ai_provider()

        settings = validate.call_args.args[0]
        self.assertEqual(provider_result, result)
        self.assertEqual("openai", settings.provider)
        self.assertEqual("gpt-5-mini", settings.model)
        self.assertEqual("saved-key", settings.api_key)

    def test_maps_capability_endpoint_uses_saved_provider_without_returning_keys(self) -> None:
        connection = self.connection_factory()
        connection.execute("UPDATE profile SET maps_provider = 'openstreetmap' WHERE id = 1")
        connection.commit()
        connection.close()
        provider_result = {
            "success": True,
            "provider": "openstreetmap",
            "provider_label": "OpenStreetMap Nominatim",
            "message": "OpenStreetMap Nominatim is ready for headquarters lookups.",
            "attribution": "© OpenStreetMap contributors",
        }
        with (
            patch.object(app_module, "get_db_connection", side_effect=self.connection_factory),
            patch.object(app_module, "validate_maps_provider", return_value=provider_result) as validate,
        ):
            result = app_module.validate_selected_maps_provider()

        settings = validate.call_args.args[0]
        self.assertEqual("openstreetmap", settings.provider)
        self.assertEqual("", settings.api_key)
        self.assertEqual(provider_result, result)


class LocalBrowserBoundaryTests(unittest.TestCase):
    def setUp(self) -> None:
        self.client = TestClient(app_module.app, base_url="http://127.0.0.1:8001")

    def tearDown(self) -> None:
        self.client.close()

    def test_non_loopback_host_is_rejected_before_routing(self) -> None:
        response = self.client.get("/api/version", headers={"Host": "attacker.example"})
        self.assertEqual(400, response.status_code)
        self.assertEqual("Invalid local Host header.", response.json()["detail"])

    def test_cross_site_writes_are_rejected_but_local_clients_reach_routing(self) -> None:
        hostile_origin = self.client.post(
            "/api/not-a-real-route",
            headers={"Origin": "https://attacker.example"},
        )
        hostile_referer = self.client.post(
            "/api/not-a-real-route",
            headers={"Referer": "https://attacker.example/form"},
        )
        cross_site_fetch = self.client.post(
            "/api/not-a-real-route",
            headers={"Sec-Fetch-Site": "same-site"},
        )
        same_origin = self.client.post(
            "/api/not-a-real-route",
            headers={"Origin": "http://127.0.0.1:8001", "Sec-Fetch-Site": "same-origin"},
        )
        local_cli = self.client.post("/api/not-a-real-route")
        self.assertEqual(403, hostile_origin.status_code)
        self.assertEqual(403, hostile_referer.status_code)
        self.assertEqual(403, cross_site_fetch.status_code)
        self.assertEqual(404, same_origin.status_code)
        self.assertEqual(404, local_cli.status_code)

    def test_oversized_resume_is_rejected_without_changing_profile(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "profile.db"
            connection = sqlite3.connect(database_path)
            connection.executescript("""
                CREATE TABLE profile (
                    id INTEGER PRIMARY KEY, base_resume_text TEXT, suggested_keywords TEXT
                );
                INSERT INTO profile VALUES (1, 'Original resume', '');
            """)
            connection.close()

            def connection_factory():
                test_connection = sqlite3.connect(database_path)
                test_connection.row_factory = sqlite3.Row
                return test_connection

            with patch.object(app_module, "get_db_connection", connection_factory):
                response = self.client.post(
                    "/api/profile/upload-resume",
                    headers={"Origin": "http://127.0.0.1:8001"},
                    files={
                        "file": (
                            "resume.txt",
                            b"a" * (app_module.MAX_RESUME_UPLOAD_BYTES + 1),
                            "text/plain",
                        )
                    },
                )

            self.assertEqual(413, response.status_code)
            connection = sqlite3.connect(database_path)
            stored = connection.execute("SELECT base_resume_text FROM profile WHERE id = 1").fetchone()[0]
            connection.close()
            self.assertEqual("Original resume", stored)

    def test_uploaded_resume_is_returned_without_being_saved_before_confirmation(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "profile.db"
            connection = sqlite3.connect(database_path)
            connection.executescript("""
                CREATE TABLE profile (
                    id INTEGER PRIMARY KEY, base_resume_text TEXT, suggested_keywords TEXT
                );
                INSERT INTO profile VALUES (1, 'Original resume', 'existing');
            """)
            connection.close()

            def connection_factory():
                test_connection = sqlite3.connect(database_path)
                test_connection.row_factory = sqlite3.Row
                return test_connection

            with patch.object(app_module, "get_db_connection", connection_factory):
                response = self.client.post(
                    "/api/profile/upload-resume",
                    headers={"Origin": "http://127.0.0.1:8001"},
                    files={"file": ("resume.md", b"# Imported resume", "text/markdown")},
                )

            self.assertEqual(200, response.status_code)
            self.assertEqual("# Imported resume", response.json()["resume_text"])
            connection = sqlite3.connect(database_path)
            stored = connection.execute(
                "SELECT base_resume_text, suggested_keywords FROM profile WHERE id = 1"
            ).fetchone()
            connection.close()
            self.assertEqual(("Original resume", "existing"), stored)


class ManualJobImportTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.database_path = Path(self.temp_dir.name) / "manual-import.db"
        connection = sqlite3.connect(self.database_path)
        connection.executescript("""
            CREATE TABLE profile (
                id INTEGER PRIMARY KEY, base_resume_text TEXT, gemini_api_key TEXT
            );
            CREATE TABLE jobs (
                id INTEGER PRIMARY KEY AUTOINCREMENT, title TEXT, company TEXT,
                description TEXT, url TEXT UNIQUE, match_score INTEGER,
                match_analysis TEXT, date_found TEXT, status TEXT, location TEXT,
                work_arrangement TEXT, employment_type TEXT, compensation TEXT,
                source TEXT, last_checked_at TEXT, is_expired INTEGER,
                expiration_reason TEXT
            );
            CREATE TABLE applications (
                id INTEGER PRIMARY KEY AUTOINCREMENT, job_id INTEGER, status TEXT
            );
            CREATE TABLE job_suppressions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                url_fingerprint TEXT NOT NULL UNIQUE, hostname TEXT NOT NULL,
                company TEXT NOT NULL DEFAULT '', title TEXT NOT NULL DEFAULT '',
                deleted_at TEXT NOT NULL, deletion_source TEXT NOT NULL
            );
            INSERT INTO profile VALUES (1, 'Experienced data leader', 'test-key');
        """)
        connection.close()

    def tearDown(self) -> None:
        self.temp_dir.cleanup()

    def connection_factory(self):
        connection = sqlite3.connect(self.database_path)
        connection.row_factory = sqlite3.Row
        return connection

    def test_preview_rejects_private_network_url_before_database_or_browser(self) -> None:
        database = unittest.mock.Mock()
        scraper = unittest.mock.Mock()
        with (
            patch.object(app_module, "get_db_connection", database),
            patch.object(app_module, "inspect_job_posting", scraper),
            self.assertRaisesRegex(app_module.HTTPException, "Local, private-network"),
        ):
            app_module.preview_manual_job(
                app_module.ManualJobPreviewRequest(url="http://127.0.0.1/jobs/private")
            )
        database.assert_not_called()
        scraper.assert_not_called()

    def test_preview_extracts_editable_fields_without_saving(self) -> None:
        extracted = {
            "title": "Data Director", "company": "Example Health",
            "description": "Lead enterprise data strategy and analytics. " * 3,
            "location": "Remote", "work_arrangement": "remote",
            "employment_type": "full_time", "compensation": "$150,000 per year",
        }
        with (
            patch.object(app_module, "get_db_connection", side_effect=self.connection_factory),
            patch.object(app_module, "validate_public_http_url", return_value=(True, "")),
            patch.object(app_module, "inspect_job_posting", return_value={
                "status": "ok", "reason": "", "details": extracted,
            }),
        ):
            result = app_module.preview_manual_job(
                app_module.ManualJobPreviewRequest(url="https://example.test/jobs/123?utm_source=email")
            )

        self.assertFalse(result["duplicate"])
        self.assertTrue(result["extraction_succeeded"])
        self.assertEqual("https://example.test/jobs/123", result["job"]["url"])
        self.assertEqual("Data Director", result["job"]["title"])
        connection = self.connection_factory()
        self.assertEqual(0, connection.execute("SELECT COUNT(*) FROM jobs").fetchone()[0])
        connection.close()

    def test_save_scores_and_inserts_reviewed_job(self) -> None:
        request = app_module.ManualJobSaveRequest(
            url="https://example.test/jobs/123", title="Data Director",
            company="Example Health",
            description="Lead enterprise data strategy and analytics. " * 3,
            location="Remote", work_arrangement="remote", employment_type="full_time",
            compensation="$150,000 per year",
        )
        with (
            patch.object(app_module, "get_db_connection", side_effect=self.connection_factory),
            patch.object(app_module, "validate_public_http_url", return_value=(True, "")),
            patch.object(app_module, "analyze_job_match", return_value={
                "success": True, "match_score": 88, "match_analysis": "Strong match."
            }),
        ):
            result = app_module.save_manual_job(request)

        self.assertEqual(88, result["match_score"])
        connection = self.connection_factory()
        stored = connection.execute("SELECT * FROM jobs WHERE id = ?", (result["job_id"],)).fetchone()
        connection.close()
        self.assertEqual("https://example.test/jobs/123", stored["url"])
        self.assertEqual("example.test", stored["source"])
        self.assertEqual("matched", stored["status"])

    def test_stopped_post_save_analysis_preserves_unscored_job(self) -> None:
        operation_id = str(uuid4())
        request = app_module.ManualJobSaveRequest(
            url="https://example.test/jobs/stop-analysis", title="Data Director",
            company="Example Health",
            description="Lead enterprise data strategy and analytics. " * 3,
        )

        def stop_during_match(*args, cancel_check=None, **kwargs):
            self.assertTrue(app_module.cancel_operation(operation_id)["active"])
            cancel_check()

        with (
            patch.object(app_module, "get_db_connection", side_effect=self.connection_factory),
            patch.object(app_module, "validate_public_http_url", return_value=(True, "")),
            patch.object(app_module, "analyze_job_match", side_effect=stop_during_match),
        ):
            result = app_module.save_manual_job(request, operation_id)

        self.assertTrue(result["success"])
        self.assertTrue(result["cancelled"])
        self.assertTrue(result["saved"])
        connection = self.connection_factory()
        stored = connection.execute(
            "SELECT match_score, match_analysis FROM jobs WHERE id = ?", (result["job_id"],)
        ).fetchone()
        connection.close()
        self.assertIsNone(stored["match_score"])
        self.assertIn("not been completed", stored["match_analysis"])

    def test_duplicate_is_reported_before_scraping(self) -> None:
        connection = self.connection_factory()
        connection.execute(
            "INSERT INTO jobs (title, company, description, url, status) VALUES (?, ?, ?, ?, ?)",
            ("Existing", "Example", "Existing description", "https://example.test/jobs/123", "matched"),
        )
        connection.commit()
        connection.close()
        scraper = unittest.mock.Mock()
        with (
            patch.object(app_module, "get_db_connection", side_effect=self.connection_factory),
            patch.object(app_module, "validate_public_http_url", return_value=(True, "")),
            patch.object(app_module, "inspect_job_posting", scraper),
        ):
            result = app_module.preview_manual_job(
                app_module.ManualJobPreviewRequest(url="https://example.test/jobs/123?utm_campaign=test")
            )
        self.assertTrue(result["duplicate"])
        scraper.assert_not_called()

    def test_suppressed_posting_is_reported_before_scraping(self) -> None:
        connection = self.connection_factory()
        record_job_suppression(
            connection,
            url="https://example.test/jobs/123?utm_source=old",
            company="Example", title="Deleted role",
            deleted_at="2026-08-08T10:00:00", deletion_source="manual",
        )
        connection.commit()
        connection.close()
        scraper = unittest.mock.Mock()
        with (
            patch.object(app_module, "get_db_connection", side_effect=self.connection_factory),
            patch.object(app_module, "validate_public_http_url", return_value=(True, "")),
            patch.object(app_module, "inspect_job_posting", scraper),
        ):
            result = app_module.preview_manual_job(
                app_module.ManualJobPreviewRequest(url="https://example.test/jobs/123?utm_campaign=new")
            )
        self.assertTrue(result["suppressed"])
        self.assertIn("Clear its suppression", result["message"])
        scraper.assert_not_called()

    def test_manual_delete_records_suppression_without_full_url(self) -> None:
        connection = self.connection_factory()
        connection.execute(
            "INSERT INTO jobs (title, company, description, url, status) VALUES (?, ?, ?, ?, ?)",
            ("Deleted role", "Example", "Description", "https://example.test/jobs/secret-123", "matched"),
        )
        job_id = connection.execute("SELECT id FROM jobs").fetchone()[0]
        connection.commit()
        connection.close()
        with patch.object(app_module, "get_db_connection", side_effect=self.connection_factory):
            result = app_module.delete_job(job_id)

        self.assertTrue(result["success"])
        connection = self.connection_factory()
        suppression = connection.execute("SELECT * FROM job_suppressions").fetchone()
        self.assertEqual("example.test", suppression["hostname"])
        self.assertEqual("manual", suppression["deletion_source"])
        self.assertNotIn("secret-123", " ".join(str(value) for value in suppression))
        self.assertEqual(0, connection.execute("SELECT COUNT(*) FROM jobs").fetchone()[0])
        connection.close()

    def test_suppressions_can_be_reviewed_and_cleared(self) -> None:
        connection = self.connection_factory()
        record_job_suppression(
            connection,
            url="https://example.test/jobs/allow-again",
            company="Example", title="Allow again",
            deleted_at="2026-08-08T10:00:00", deletion_source="manual",
        )
        connection.commit()
        connection.close()
        with patch.object(app_module, "get_db_connection", side_effect=self.connection_factory):
            review = app_module.get_job_suppressions()
            cleared = app_module.clear_job_suppression(review["items"][0]["id"])
            after = app_module.get_job_suppressions()

        self.assertEqual(1, review["count"])
        self.assertNotIn("url", review["items"][0])
        self.assertNotIn("url_fingerprint", review["items"][0])
        self.assertEqual(1, cleared["cleared"])
        self.assertEqual(0, after["count"])

    def test_preview_uses_resolved_url_and_detects_redirect_duplicate(self) -> None:
        connection = self.connection_factory()
        connection.execute(
            "INSERT INTO jobs (title, company, description, url, status) VALUES (?, ?, ?, ?, ?)",
            ("Existing", "Example", "Existing description", "https://careers.example.test/jobs/123", "matched"),
        )
        connection.commit()
        connection.close()
        with (
            patch.object(app_module, "get_db_connection", side_effect=self.connection_factory),
            patch.object(app_module, "validate_public_http_url", return_value=(True, "")),
            patch.object(app_module, "inspect_job_posting", return_value={
                "status": "partial", "reason": "", "details": {
                    "url": "https://careers.example.test/jobs/123",
                    "title": "Data Director",
                },
            }),
        ):
            result = app_module.preview_manual_job(
                app_module.ManualJobPreviewRequest(url="https://example.test/apply/123")
            )

        self.assertTrue(result["duplicate"])
        self.assertEqual("https://careers.example.test/jobs/123", result["job"]["url"])

    def test_preview_explains_access_challenge_and_preserves_manual_entry(self) -> None:
        with (
            patch.object(app_module, "get_db_connection", side_effect=self.connection_factory),
            patch.object(app_module, "validate_public_http_url", return_value=(True, "")),
            patch.object(app_module, "inspect_job_posting", return_value={
                "status": "access_challenge", "reason": "Automated access blocked.", "details": {},
            }),
        ):
            result = app_module.preview_manual_job(
                app_module.ManualJobPreviewRequest(url="https://jobs.example.test/opening/123")
            )

        self.assertFalse(result["extraction_succeeded"])
        self.assertEqual("access_challenge", result["extraction_status"])
        self.assertIn("Enter the posting details manually", result["message"])

    def test_save_without_ai_key_keeps_job_visible_as_unscored(self) -> None:
        connection = self.connection_factory()
        connection.execute("UPDATE profile SET gemini_api_key = '' WHERE id = 1")
        connection.commit()
        connection.close()
        matcher = unittest.mock.Mock()
        request = app_module.ManualJobSaveRequest(
            url="https://example.test/jobs/unscored", title="Operations Director",
            company="Example Manufacturing",
            description="Lead operations, safety, quality, and continuous improvement. " * 2,
        )
        with (
            patch.object(app_module, "get_db_connection", side_effect=self.connection_factory),
            patch.object(app_module, "validate_public_http_url", return_value=(True, "")),
            patch.object(app_module, "analyze_job_match", matcher),
        ):
            result = app_module.save_manual_job(request)

        self.assertIsNone(result["match_score"])
        matcher.assert_not_called()
        connection = self.connection_factory()
        stored_score = connection.execute(
            "SELECT match_score FROM jobs WHERE id = ?", (result["job_id"],)
        ).fetchone()[0]
        connection.close()
        self.assertIsNone(stored_score)

    def test_malformed_ai_result_does_not_prevent_save(self) -> None:
        request = app_module.ManualJobSaveRequest(
            url="https://example.test/jobs/malformed-ai", title="Operations Director",
            company="Example Manufacturing",
            description="Lead operations, safety, quality, and continuous improvement. " * 2,
        )
        with (
            patch.object(app_module, "get_db_connection", side_effect=self.connection_factory),
            patch.object(app_module, "validate_public_http_url", return_value=(True, "")),
            patch.object(app_module, "analyze_job_match", return_value=None),
        ):
            result = app_module.save_manual_job(request)

        self.assertIsNone(result["match_score"])
        connection = self.connection_factory()
        stored = connection.execute(
            "SELECT id, match_score FROM jobs WHERE id = ?", (result["job_id"],)
        ).fetchone()
        connection.close()
        self.assertIsNotNone(stored)
        self.assertIsNone(stored["match_score"])

    def test_verification_error_does_not_close_active_job(self) -> None:
        connection = self.connection_factory()
        cursor = connection.execute(
            "INSERT INTO jobs (title, company, description, url, status, is_expired) VALUES (?, ?, ?, ?, ?, 0)",
            ("Data Director", "Example", "Complete description", "https://example.test/jobs/verify", "matched"),
        )
        connection.commit()
        job_id = cursor.lastrowid
        connection.close()
        with (
            patch.object(app_module, "get_db_connection", side_effect=self.connection_factory),
            patch.object(app_module, "inspect_job_posting", return_value={
                "status": "format_drift", "details": {}, "reason": "Unexpected format.",
            }),
            self.assertRaisesRegex(app_module.HTTPException, "No job status was changed"),
        ):
            app_module.verify_job_posting(job_id)

        connection = self.connection_factory()
        stored = connection.execute(
            "SELECT status, is_expired FROM jobs WHERE id = ?", (job_id,)
        ).fetchone()
        connection.close()
        self.assertEqual("matched", stored["status"])
        self.assertEqual(0, stored["is_expired"])


class BaseResumeVersionTests(unittest.TestCase):
    def setUp(self) -> None:
        self.connection = sqlite3.connect(":memory:")
        self.connection.row_factory = sqlite3.Row
        self.connection.execute("PRAGMA foreign_keys = ON")
        self.connection.executescript("""
            CREATE TABLE profile (
                id INTEGER PRIMARY KEY, base_resume_text TEXT, resume_mode TEXT,
                active_base_resume_id INTEGER, suggested_keywords TEXT
            );
            CREATE TABLE base_resumes (
                id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT NOT NULL,
                resume_mode TEXT NOT NULL, content TEXT NOT NULL,
                evidence_json TEXT NOT NULL DEFAULT '{}',
                created_at TEXT NOT NULL, updated_at TEXT NOT NULL
            );
            CREATE TABLE base_resume_versions (
                id INTEGER PRIMARY KEY AUTOINCREMENT, base_resume_id INTEGER NOT NULL,
                version_number INTEGER NOT NULL, name TEXT NOT NULL,
                resume_mode TEXT NOT NULL, content TEXT NOT NULL,
                evidence_json TEXT NOT NULL DEFAULT '{}', created_at TEXT NOT NULL,
                UNIQUE(base_resume_id, version_number),
                FOREIGN KEY (base_resume_id) REFERENCES base_resumes(id) ON DELETE CASCADE
            );
            INSERT INTO profile VALUES (1, '', 'general_professional', NULL, 'stale');
        """)

    def tearDown(self) -> None:
        self.connection.close()

    def test_save_snapshots_changes_without_duplicate_versions(self) -> None:
        created = save_base_resume(
            self.connection, None, "Data Leadership", "technical_executive", "Version one"
        )
        activate_base_resume(self.connection, created["id"])
        unchanged = save_base_resume(
            self.connection, created["id"], "Data Leadership", "technical_executive", "Version one",
            {"skills": "", "work_samples": ""},
        )
        updated = save_base_resume(
            self.connection, created["id"], "Data Leadership", "technical_executive", "Version two"
        )

        self.assertEqual(1, created["version_number"])
        self.assertFalse(unchanged["version_created"])
        self.assertEqual(1, unchanged["version_number"])
        self.assertEqual(2, updated["version_number"])
        self.assertEqual([2, 1], [row["version_number"] for row in list_versions(self.connection, created["id"])])

    def test_restore_creates_a_new_version_and_syncs_the_active_profile(self) -> None:
        created = save_base_resume(
            self.connection, None, "Primary", "it", "Original",
            {"skills": "Python", "portfolio": "Architecture — https://example.test/work"},
        )
        activate_base_resume(self.connection, created["id"])
        save_base_resume(
            self.connection, created["id"], "Primary", "it", "Edited",
            {"skills": "Python\nGo", "certifications": "Cloud certification"},
        )

        restored = restore_version(self.connection, created["id"], 1)
        profile = self.connection.execute(
            "SELECT base_resume_text, resume_mode, suggested_keywords FROM profile WHERE id = 1"
        ).fetchone()

        self.assertEqual(3, restored["version_number"])
        self.assertEqual("Original", get_base_resume(self.connection, created["id"])["content"])
        self.assertEqual("Python", restored["professional_evidence"]["skills"])
        self.assertIn("https://example.test/work", restored["professional_evidence"]["portfolio"])
        self.assertEqual(("Original", "it", ""), tuple(profile))

    def test_evidence_only_change_creates_a_version(self) -> None:
        created = save_base_resume(
            self.connection, None, "Operations", "trades_operations", "Resume",
            {"licenses": "CDL-A"},
        )
        updated = save_base_resume(
            self.connection, created["id"], "Operations", "trades_operations", "Resume",
            {"licenses": "CDL-A", "work_samples": "Safety case study"},
        )

        self.assertTrue(updated["version_created"])
        self.assertEqual(2, updated["version_number"])
        self.assertEqual("Safety case study", updated["professional_evidence"]["work_samples"])
        self.assertEqual(2, list_versions(self.connection, created["id"])[0]["evidence_section_count"])
        backward_compatible = save_base_resume(
            self.connection, created["id"], "Operations", "trades_operations", "Updated resume"
        )
        self.assertEqual("CDL-A", backward_compatible["professional_evidence"]["licenses"])
        self.assertEqual("Safety case study", backward_compatible["professional_evidence"]["work_samples"])
        self.assertNotIn("evidence_json", backward_compatible)

    def test_deleting_active_resume_selects_fallback_but_protects_last_resume(self) -> None:
        first = save_base_resume(self.connection, None, "First", "it", "One")
        second = save_base_resume(self.connection, None, "Second", "general_professional", "Two")
        activate_base_resume(self.connection, first["id"])

        replacement = delete_base_resume(self.connection, first["id"])
        self.assertEqual(second["id"], replacement["id"])
        self.assertEqual(second["id"], self.connection.execute(
            "SELECT active_base_resume_id FROM profile WHERE id = 1"
        ).fetchone()[0])
        with self.assertRaises(LastBaseResumeError):
            delete_base_resume(self.connection, second["id"])


class HeadquartersPreferenceTests(unittest.TestCase):
    def test_profile_save_persists_disabled_us_preference(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            db_path = Path(temp_dir) / "profile.db"
            connection = sqlite3.connect(db_path)
            connection.executescript("""
                CREATE TABLE profile (
                    id INTEGER PRIMARY KEY, name TEXT, email TEXT, phone TEXT,
                    github TEXT, linkedin TEXT, website TEXT, base_resume_text TEXT,
                    resume_mode TEXT, ai_provider TEXT, ai_model TEXT,
                    maps_provider TEXT,
                    prefer_us_headquarters INTEGER,
                    suggested_keywords TEXT, active_base_resume_id INTEGER
                );
                CREATE TABLE base_resumes (
                    id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT, resume_mode TEXT,
                    content TEXT, evidence_json TEXT NOT NULL DEFAULT '{}',
                    created_at TEXT, updated_at TEXT
                );
                CREATE TABLE base_resume_versions (
                    id INTEGER PRIMARY KEY AUTOINCREMENT, base_resume_id INTEGER,
                    version_number INTEGER, name TEXT, resume_mode TEXT,
                    content TEXT, evidence_json TEXT NOT NULL DEFAULT '{}', created_at TEXT
                );
                INSERT INTO profile VALUES (1, '', '', '', '', '', '', '', '', 'gemini', 'gemini-2.5-flash', 'google', 1, '', NULL);
            """)
            connection.close()

            def open_test_database():
                test_connection = sqlite3.connect(db_path)
                test_connection.row_factory = sqlite3.Row
                return test_connection

            profile = app_module.ProfileUpdate(
                name="Candidate", email="candidate@example.test", phone="",
                github="", linkedin="", website="", base_resume_text="Resume",
                resume_mode="general_professional", prefer_us_headquarters=False,
            )
            with patch.object(app_module, "get_db_connection", side_effect=open_test_database):
                app_module.update_profile(profile)

            connection = sqlite3.connect(db_path)
            stored = connection.execute(
                "SELECT prefer_us_headquarters FROM profile WHERE id = 1"
            ).fetchone()[0]
            connection.close()
            self.assertEqual(0, stored)

    def test_us_preference_changes_places_query(self) -> None:
        self.assertEqual("Example Co United States headquarters", _headquarters_query("Example Co", True))
        self.assertEqual("Example Co global headquarters", _headquarters_query("Example Co", False))

    def test_us_address_requires_explicit_country(self) -> None:
        self.assertTrue(_looks_like_us_address("New York, NY 10001, USA"))
        self.assertTrue(_looks_like_us_address("Chicago, IL, United States"))
        self.assertFalse(_looks_like_us_address("Chennai, Tamil Nadu 600100, India"))

    @patch("urllib.request.urlopen")
    def test_us_preference_rejects_non_us_places_result(self, urlopen) -> None:
        us_response = unittest.mock.MagicMock()
        us_response.__enter__.return_value.read.return_value = (
            b'{"places":[{"formattedAddress":"Chennai, Tamil Nadu 600100, India",'
            b'"addressComponents":[{"shortText":"IN","types":["country"]}]}]}'
        )
        global_response = unittest.mock.MagicMock()
        global_response.__enter__.return_value.read.return_value = b'{"places":[]}'
        urlopen.side_effect = [us_response, global_response]
        with patch("utils.get_gemini_api_key", return_value=""):
            self.assertEqual(
                "Unknown",
                find_us_headquarters("Future Works", api_key="", google_maps_key="maps-key", prefer_us=True),
            )
        self.assertEqual(2, urlopen.call_count)

    @patch("urllib.request.urlopen")
    def test_global_preference_accepts_formatted_global_result(self, urlopen) -> None:
        response = unittest.mock.MagicMock()
        response.__enter__.return_value.read.return_value = (
            b'{"places":[{"formattedAddress":"London EC1V 3AG, United Kingdom",'
            b'"addressComponents":[{"shortText":"GB","types":["country"]}]}]}'
        )
        urlopen.return_value = response
        self.assertEqual(
            "London EC1V 3AG, United Kingdom",
            find_us_headquarters("Example Co", api_key="", google_maps_key="maps-key", prefer_us=False),
        )


class SearchQualityTests(unittest.TestCase):
    class FakeHttpResponse:
        def __init__(self, payload: bytes):
            self.payload = payload
            self.read_limit = None

        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

        def read(self, limit):
            self.read_limit = limit
            return self.payload

    @staticmethod
    def provider_diagnostics(provider="lever"):
        return {
            provider: {
                "raw_candidates": 0,
                "valid_discovered": 0,
                "new_candidates": 0,
                "accepted": 0,
                "errors": [],
                "candidate_budget_exhausted": False,
                "partial_results": {
                    "oversized_responses": 0,
                    "candidate_limit_hits": 0,
                    "timeouts": 0,
                },
                "rejection_reasons": {},
            }
        }

    def test_generic_smartrecruiters_career_pages_are_not_job_postings(self) -> None:
        self.assertFalse(is_specific_job_url("https://careers.smartrecruiters.com/QADInc/corporate-careers"))
        self.assertTrue(is_specific_job_url("https://jobs.smartrecruiters.com/Example/12345-data-analyst"))

    def test_tracking_variants_share_one_canonical_url(self) -> None:
        base = "https://jobs.lever.co/example/abc-123"
        self.assertEqual(base, canonicalize_job_url(base + "/?source=linkedin#apply"))

    def test_search_skips_suppressed_canonical_url_before_scraping(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "suppressed-search.db"
            connection = sqlite3.connect(database_path)
            connection.executescript("""
                CREATE TABLE profile (
                    id INTEGER PRIMARY KEY, base_resume_text TEXT,
                    gemini_api_key TEXT, suggested_keywords TEXT
                );
                CREATE TABLE jobs (
                    id INTEGER PRIMARY KEY, url TEXT, status TEXT
                );
                CREATE TABLE job_suppressions (
                    id INTEGER PRIMARY KEY, url_fingerprint TEXT UNIQUE,
                    hostname TEXT, company TEXT, title TEXT,
                    deleted_at TEXT, deletion_source TEXT
                );
                INSERT INTO profile VALUES (1, 'Experienced data leader', 'test-key', '');
            """)
            record_job_suppression(
                connection,
                url="https://jobs.lever.co/example/abc-123?source=old",
                company="Example", title="Data leader",
                deleted_at="2026-08-08T10:00:00", deletion_source="manual",
            )
            connection.commit()
            connection.close()

            def connection_factory():
                test_connection = sqlite3.connect(database_path)
                test_connection.row_factory = sqlite3.Row
                return test_connection

            scraper = unittest.mock.Mock()
            matcher = unittest.mock.Mock()
            with (
                patch.object(searcher_module, "get_db_connection", side_effect=connection_factory),
                patch.object(searcher_module, "search_yahoo_jobs", return_value=[{
                    "url": "https://jobs.lever.co/example/abc-123?utm_source=new"
                }]),
                patch.object(searcher_module, "inspect_job_posting", scraper),
                patch.object(searcher_module, "analyze_job_matches_batch", matcher),
            ):
                result = searcher_module.run_job_search_and_matching("data leader")

        self.assertEqual(0, result["jobs_added"])
        self.assertEqual(1, result["provider_health"]["lever"]["skipped_suppressed"])
        scraper.assert_not_called()
        matcher.assert_not_called()

    def test_url_fingerprint_ignores_tracking_variants(self) -> None:
        self.assertEqual(
            job_url_fingerprint("https://jobs.lever.co/example/abc-123?source=linkedin"),
            job_url_fingerprint("https://jobs.lever.co/example/abc-123?utm_campaign=email"),
        )

    def test_ashby_application_route_canonicalizes_to_the_posting(self) -> None:
        posting = "https://jobs.ashbyhq.com/change/c73f61f9-e17d-4d29-bc58-40702ef50ef2"

        self.assertEqual(posting, canonicalize_job_url(f"{posting}/application"))
        self.assertEqual(posting, canonicalize_job_url(f"{posting}/application/?utm_source=email"))

    def test_unknown_site_keeps_functional_query_but_removes_tracking(self) -> None:
        self.assertEqual(
            "https://careers.example.test/opening?job=123",
            canonicalize_job_url(
                "https://careers.example.test/opening?utm_source=email&job=123#apply"
            ),
        )

    def test_public_url_validation_blocks_local_and_private_addresses(self) -> None:
        for url in (
            "http://127.0.0.1/job", "http://localhost/job",
            "http://10.10.1.4/job", "http://169.254.169.254/latest/meta-data",
            "file:///etc/passwd",
        ):
            with self.subTest(url=url):
                self.assertFalse(is_public_http_url(url))

    @patch("searcher.socket.getaddrinfo")
    def test_public_url_validation_requires_only_global_dns_results(self, getaddrinfo) -> None:
        getaddrinfo.return_value = [
            (2, 1, 6, "", ("93.184.216.34", 443)),
        ]
        self.assertTrue(is_public_http_url("https://jobs.example.test/opening"))
        getaddrinfo.return_value = [
            (2, 1, 6, "", ("93.184.216.34", 443)),
            (2, 1, 6, "", ("127.0.0.1", 443)),
        ]
        self.assertFalse(is_public_http_url("https://mixed.example.test/opening"))

    @patch("searcher.socket.getaddrinfo")
    def test_browser_route_guard_blocks_private_redirects_and_subresources(self, getaddrinfo) -> None:
        getaddrinfo.return_value = [(2, 1, 6, "", ("93.184.216.34", 443))]

        class FakeContext:
            handler = None

            def route(self, pattern, handler):
                self.handler = handler

        class FakeRoute:
            continued = False
            aborted = False

            def continue_(self):
                self.continued = True

            def abort(self, reason):
                self.aborted = reason == "blockedbyclient"

        class FakeRequest:
            def __init__(self, url):
                self.url = url

        context = FakeContext()
        _protect_browser_network(context)
        public_route = FakeRoute()
        private_route = FakeRoute()
        context.handler(public_route, FakeRequest("https://assets.example.test/app.js"))
        context.handler(private_route, FakeRequest("http://127.0.0.1/private"))
        self.assertTrue(public_route.continued)
        self.assertTrue(private_route.aborted)

    def test_browser_error_page_is_not_useful_job_content(self) -> None:
        self.assertFalse(is_useful_job_details({
            "title": "Sorry, Internet Explorer 11 is no longer supported by SmartRecruiters",
            "company": "LegalAndGeneral Enterprise Data Architect",
            "description": "Please update your browser. " * 10,
        }))

    def test_real_job_content_is_useful(self) -> None:
        self.assertTrue(is_useful_job_details({
            "title": "Chief Data Officer",
            "company": "Example Health",
            "description": "Lead enterprise data strategy, governance, analytics, and a multidisciplinary team. " * 3,
        }))

    def test_provider_detection_and_metadata_normalization(self) -> None:
        self.assertEqual("greenhouse", provider_for_url("https://job-boards.greenhouse.io/acme/jobs/123"))
        self.assertEqual("ashby", provider_for_url("https://jobs.ashbyhq.com/acme/123"))
        self.assertEqual("lever", provider_for_url("https://jobs.eu.lever.co/acme/123"))
        metadata = _metadata_from_text("Remote full time role. Compensation $120,000 - $150,000 per year.")
        self.assertEqual("remote", metadata["work_arrangement"])
        self.assertEqual("full_time", metadata["employment_type"])
        self.assertIn("$120,000", metadata["compensation"])

    def test_json_ld_location_normalizes_structured_country_values(self) -> None:
        posting = {
            "jobLocation": [
                {
                    "@type": "Place",
                    "address": {
                        "@type": "PostalAddress",
                        "addressLocality": "Los Angeles, CA,US",
                        "addressRegion": "CA",
                        "addressCountry": {"@type": "Country", "name": "US"},
                    },
                },
                {
                    "@type": "Place",
                    "address": {
                        "addressLocality": "Dayton",
                        "addressRegion": "OH",
                        "addressCountry": {"@type": "Country", "name": "US"},
                    },
                },
            ]
        }

        self.assertEqual(
            "Los Angeles, CA, US; Dayton, OH, US",
            _job_location_from_json_ld(posting),
        )

    def test_popular_job_board_structured_metadata_shapes_are_normalized(self) -> None:
        samples = {
            "ziprecruiter": ({
                "employmentType": "PART_TIME",
                "baseSalary": {
                    "@type": "MonetaryAmount", "currency": "USD",
                    "value": {"@type": "QuantitativeValue", "minValue": 18, "maxValue": 20, "unitText": "HOUR"},
                },
            }, {"work_arrangement": "", "employment_type": "part_time", "compensation": "USD 18–20 per hour"}),
            "glassdoor": ({
                "employmentType": ["FULL_TIME"], "jobLocationType": "TELECOMMUTE",
                "estimatedSalary": {
                    "@type": "MonetaryAmount", "currency": "USD",
                    "value": {"@type": "QuantitativeValue", "minValue": 77447, "maxValue": 113374, "unitText": "YEAR"},
                },
            }, {"work_arrangement": "remote", "employment_type": "full_time", "compensation": "USD 77,447–113,374 per year"}),
            "dice": ({
                "employmentType": "CONTRACTOR", "jobLocationType": "TELECOMMUTE",
                "baseSalary": {"@type": "MonetaryAmount", "currency": "USD", "minValue": 70, "maxValue": 85},
            }, {"work_arrangement": "remote", "employment_type": "contract", "compensation": "USD 70–85"}),
        }

        for platform, (posting, expected) in samples.items():
            with self.subTest(platform=platform):
                self.assertEqual(expected, _structured_job_metadata(posting))

    def test_remote_applicant_location_is_used_when_job_location_is_absent(self) -> None:
        posting = {
            "jobLocationType": "TELECOMMUTE",
            "applicantLocationRequirements": [
                {"@type": "Country", "name": "USA"},
                {"@type": "AdministrativeArea", "name": "District of Columbia"},
            ],
        }

        self.assertEqual("USA; District of Columbia", _job_location_from_json_ld(posting))

    def test_provider_format_drift_creates_user_alert(self) -> None:
        health = {"ashby": {"raw_candidates": 4, "valid_discovered": 4, "new_candidates": 3, "accepted": 0, "errors": []}}
        alerts = provider_alerts_from_health(health)
        self.assertEqual("content_format_drift", alerts[0]["code"])

    def test_lever_api_response_is_normalized_without_browser_scraping(self) -> None:
        payload = json.dumps({
            "id": "posting-123",
            "text": "Director of Data",
            "descriptionPlain": "Lead enterprise data strategy, governance, analytics, architecture, and delivery across the organization. " * 2,
            "categories": {
                "location": "Dayton, OH", "allLocations": ["Dayton, OH", "Remote"],
                "commitment": "Full-time",
            },
            "workplaceType": "hybrid",
            "salaryRange": {"currency": "USD", "interval": "year", "min": 150000, "max": 190000},
        }).encode("utf-8")
        with patch.object(
            searcher_module.urllib.request,
            "urlopen",
            return_value=self.FakeHttpResponse(payload),
        ) as urlopen:
            outcome = searcher_module._lever_posting_from_api(
                "https://jobs.lever.co/acme-corp/posting-123"
            )

        self.assertEqual("ok", outcome["status"])
        self.assertEqual("Director of Data", outcome["details"]["title"])
        self.assertEqual("Acme Corp", outcome["details"]["company"])
        self.assertEqual("Dayton, OH; Remote", outcome["details"]["location"])
        self.assertEqual("hybrid", outcome["details"]["work_arrangement"])
        self.assertEqual("full_time", outcome["details"]["employment_type"])
        self.assertEqual("USD 150,000–190,000 per year", outcome["details"]["compensation"])
        request = urlopen.call_args.args[0]
        self.assertEqual(
            "https://api.lever.co/v0/postings/acme-corp/posting-123",
            request.full_url,
        )

    def test_lever_removed_posting_is_stale_without_browser_fallback(self) -> None:
        stale = {"status": "stale", "details": {}, "reason": "No longer published."}
        browser = unittest.mock.Mock()
        with (
            patch.object(searcher_module, "validate_public_http_url", return_value=(True, "")),
            patch.object(searcher_module, "_lever_posting_from_api", return_value=stale),
            patch.object(searcher_module, "_scrape_job_details_browser", browser),
        ):
            outcome = searcher_module.inspect_job_posting(
                "https://jobs.lever.co/acme/posting-123"
            )

        self.assertEqual("stale", outcome["status"])
        browser.assert_not_called()

    def test_lever_invalid_successful_response_is_format_drift(self) -> None:
        payload = json.dumps({"id": "posting-123", "categories": []}).encode("utf-8")
        with patch.object(
            searcher_module.urllib.request,
            "urlopen",
            return_value=self.FakeHttpResponse(payload),
        ):
            outcome = searcher_module._lever_posting_from_api(
                "https://jobs.lever.co/acme/posting-123"
            )

        self.assertEqual("format_drift", outcome["status"])

    def test_lever_api_not_found_is_classified_as_stale(self) -> None:
        error = searcher_module.urllib.error.HTTPError(
            "https://api.lever.co/v0/postings/acme/posting-123",
            404,
            "Not Found",
            hdrs=None,
            fp=None,
        )
        with patch.object(searcher_module.urllib.request, "urlopen", side_effect=error):
            outcome = searcher_module._lever_posting_from_api(
                "https://jobs.lever.co/acme/posting-123"
            )

        self.assertEqual("stale", outcome["status"])

    def test_lever_api_unavailable_uses_browser_fallback(self) -> None:
        def browser_fallback(_url, _allow_partial, diagnostic):
            diagnostic.update({"status": "ok", "reason": ""})
            return {"title": "Fallback role", "company": "Acme", "description": "Complete description"}

        with (
            patch.object(searcher_module, "validate_public_http_url", return_value=(True, "")),
            patch.object(searcher_module, "_lever_posting_from_api", return_value={
                "status": "api_unavailable", "details": {}, "reason": "Timed out.",
            }),
            patch.object(searcher_module, "_scrape_job_details_browser", side_effect=browser_fallback),
        ):
            outcome = searcher_module.inspect_job_posting(
                "https://jobs.lever.co/acme/posting-123"
            )

        self.assertEqual("ok", outcome["status"])
        self.assertTrue(outcome["used_browser_fallback"])

    def test_lever_stale_postings_do_not_create_format_drift_alert(self) -> None:
        health = {
            "lever": {
                "raw_candidates": 3, "valid_discovered": 3, "new_candidates": 3,
                "accepted": 0, "errors": [],
                "rejection_reasons": {"stale": 3, "format_drift": 0},
            }
        }

        alerts = provider_alerts_from_health(health)

        self.assertEqual(["stale_postings"], [alert["code"] for alert in alerts])

    def test_yahoo_search_bounds_response_bytes_and_reports_partial_results(self) -> None:
        posting_url = "https://jobs.lever.co/acme/posting-123"
        response = self.FakeHttpResponse(
            f'<a href="{posting_url}">Role</a>'.encode("utf-8")
            + b"x" * searcher_module.MAX_SEARCH_RESPONSE_BYTES
        )
        diagnostics = self.provider_diagnostics()
        with (
            patch.object(searcher_module, "PROVIDER_DOMAINS", {"lever": "lever.co"}),
            patch.object(searcher_module.urllib.request, "urlopen", return_value=response) as urlopen,
        ):
            results = searcher_module.search_yahoo_jobs("data leader", diagnostics=diagnostics)

        self.assertEqual([posting_url], [result["url"] for result in results])
        self.assertEqual(searcher_module.MAX_SEARCH_RESPONSE_BYTES + 1, response.read_limit)
        self.assertEqual(searcher_module.SEARCH_HTTP_TIMEOUT_SECONDS, urlopen.call_args.kwargs["timeout"])
        self.assertEqual(1, diagnostics["lever"]["partial_results"]["oversized_responses"])

    def test_yahoo_search_caps_candidates_per_provider(self) -> None:
        links = "".join(
            f'<a href="https://jobs.lever.co/acme/posting-{index}">Role</a>'
            for index in range(searcher_module.MAX_PROVIDER_CANDIDATES + 3)
        ).encode("utf-8")
        diagnostics = self.provider_diagnostics()
        with (
            patch.object(searcher_module, "PROVIDER_DOMAINS", {"lever": "lever.co"}),
            patch.object(
                searcher_module.urllib.request,
                "urlopen",
                return_value=self.FakeHttpResponse(links),
            ),
        ):
            results = searcher_module.search_yahoo_jobs("data leader", diagnostics=diagnostics)

        self.assertEqual(searcher_module.MAX_PROVIDER_CANDIDATES, len(results))
        self.assertEqual(1, diagnostics["lever"]["partial_results"]["candidate_limit_hits"])

    def test_provider_candidate_cap_applies_across_search_queries(self) -> None:
        results = []
        seen_urls = set()
        counts = {"lever": 0}
        health = self.provider_diagnostics()

        appended = [
            searcher_module._append_candidate_with_budget(
                {"url": f"https://jobs.lever.co/acme/posting-{index}"},
                results,
                seen_urls,
                counts,
                health,
            )
            for index in range(searcher_module.MAX_PROVIDER_CANDIDATES + 2)
        ]

        self.assertEqual(searcher_module.MAX_PROVIDER_CANDIDATES, sum(appended))
        self.assertEqual(searcher_module.MAX_PROVIDER_CANDIDATES, len(results))
        self.assertTrue(health["lever"]["candidate_budget_exhausted"])

    def test_yahoo_timeout_is_a_partial_result_instead_of_an_unbounded_wait(self) -> None:
        diagnostics = self.provider_diagnostics()
        with (
            patch.object(searcher_module, "PROVIDER_DOMAINS", {"lever": "lever.co"}),
            patch.object(searcher_module.urllib.request, "urlopen", side_effect=TimeoutError),
        ):
            results = searcher_module.search_yahoo_jobs("data leader", diagnostics=diagnostics)

        self.assertEqual([], results)
        self.assertEqual(1, diagnostics["lever"]["partial_results"]["timeouts"])
        self.assertEqual([], diagnostics["lever"]["errors"])

    def test_provider_budget_alert_explains_partial_results(self) -> None:
        health = self.provider_diagnostics()
        health["lever"]["candidate_budget_exhausted"] = True
        health["lever"]["partial_results"]["oversized_responses"] = 1

        alerts = provider_alerts_from_health(health)

        self.assertEqual(["partial_results"], [alert["code"] for alert in alerts])
        self.assertIn("1 MiB", alerts[0]["message"])
        self.assertIn("25-posting", alerts[0]["message"])

    def test_job_description_is_bounded_at_the_matching_storage_boundary(self) -> None:
        original = {"title": "Role", "description": "x" * (searcher_module.MAX_JOB_DESCRIPTION_CHARS + 100)}

        bounded = searcher_module._bounded_job_details(original)

        self.assertEqual(searcher_module.MAX_JOB_DESCRIPTION_CHARS, len(bounded["description"]))
        self.assertGreater(len(original["description"]), len(bounded["description"]))

    def test_generic_extraction_prefers_structured_job_data_in_child_frame(self) -> None:
        class Scope:
            def __init__(self, url):
                self.url = url

        page = Scope("https://careers.example.test/openings/123")
        frame = Scope("https://embedded.example.test/jobs/123")
        page.main_frame = page
        page.frames = [page, frame]
        framed_posting = {"@type": "JobPosting", "title": "Framed role"}

        with (
            patch.object(
                searcher_module,
                "_job_posting_json_ld",
                side_effect=lambda scope: framed_posting if scope is frame else {},
            ),
            patch.object(searcher_module, "_largest_text", return_value=""),
        ):
            scope, posting, has_job_frame = searcher_module._select_job_content_scope(page)

        self.assertIs(frame, scope)
        self.assertEqual(framed_posting, posting)
        self.assertTrue(has_job_frame)


class ResumeTemplateTests(unittest.TestCase):
    def test_sections_are_normalized_into_mode_order(self) -> None:
        source = "# Candidate\n\n## Experience\nRole details\n\n## Skills\nPython\n\n## Summary\nLeader"
        result = apply_resume_section_template(source, "it")
        self.assertLess(result.index("## Summary"), result.index("## Technical Skills"))
        self.assertLess(result.index("## Technical Skills"), result.index("## Professional Experience"))

    def test_portfolio_and_work_samples_map_to_role_specific_section(self) -> None:
        source = "# Candidate\n\n## Portfolio\nhttps://example.test/work\n\n## Experience\nRole details"
        result = apply_resume_section_template(source, "it")
        self.assertIn("## Portfolio & Work Samples", result)
        self.assertLess(result.index("## Professional Experience"), result.index("## Portfolio & Work Samples"))

    def test_unknown_sections_are_rejected_instead_of_silently_dropped(self) -> None:
        with self.assertRaisesRegex(ValueError, "unsupported section"):
            apply_resume_section_template("# Candidate\n\n## Mystery Material\nText", "it")


class SourceDiagnosticHistoryTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.database_path = Path(self.temp_dir.name) / "source-diagnostics.db"
        connection = sqlite3.connect(self.database_path)
        connection.executescript("""
            CREATE TABLE source_diagnostics (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                recorded_at TEXT NOT NULL,
                provider TEXT NOT NULL,
                diagnostic_code TEXT NOT NULL,
                counters_json TEXT NOT NULL
            );
        """)
        connection.close()

    def tearDown(self) -> None:
        self.temp_dir.cleanup()

    def connection_factory(self):
        connection = sqlite3.connect(self.database_path)
        connection.row_factory = sqlite3.Row
        return connection

    @staticmethod
    def search_result(code="content_format_drift", recorded_value=4) -> dict:
        return {
            "provider_alerts": [{
                "provider": "lever", "code": code,
                "message": "Do not retain https://private.example/jobs/secret",
                "unexpected": "resume and credential data",
            }],
            "provider_health": {
                "lever": {
                    "raw_candidates": recorded_value,
                    "valid_discovered": -3,
                    "new_candidates": "2",
                    "accepted": 0,
                    "rejected": 2,
                    "errors": ["API key secret-value", "https://private.example"],
                    "search": "private keywords in Dayton",
                    "resume": "private resume text",
                    "rejection_reasons": {
                        "format_drift": 2,
                        "provider_error": MAX_COUNTER_VALUE + 100,
                        "raw_exception": "secret exception",
                    },
                    "partial_results": {"timeouts": 1, "private": "secret"},
                }
            },
        }

    def test_persistence_allowlists_and_bounds_diagnostic_data(self) -> None:
        connection = self.connection_factory()
        inserted = persist_source_diagnostics(
            connection, self.search_result(), "2026-08-08T12:00:00"
        )
        connection.commit()
        stored_json = connection.execute("SELECT counters_json FROM source_diagnostics").fetchone()[0]
        history = list_source_diagnostics(connection)
        connection.close()

        self.assertEqual(1, inserted)
        self.assertEqual(0, history["items"][0]["counters"]["valid_discovered"])
        self.assertEqual(MAX_COUNTER_VALUE, history["items"][0]["counters"]["provider_error"])
        self.assertEqual(1, history["items"][0]["counters"]["timeouts"])
        self.assertEqual(2, history["items"][0]["counters"]["search_errors"])
        for private_value in ("private.example", "secret-value", "private keywords", "resume text", "secret exception"):
            self.assertNotIn(private_value, stored_json)

    def test_unknown_provider_and_code_are_not_persisted(self) -> None:
        result = self.search_result()
        result["provider_alerts"].extend([
            {"provider": "unknown", "code": "provider_error"},
            {"provider": "lever", "code": "raw_exception"},
        ])
        connection = self.connection_factory()
        inserted = persist_source_diagnostics(connection, result, "2026-08-08T12:00:00")
        connection.commit()
        count = connection.execute("SELECT COUNT(*) FROM source_diagnostics").fetchone()[0]
        connection.close()
        self.assertEqual(1, inserted)
        self.assertEqual(1, count)

    def test_informational_search_note_remains_retrievable_after_dismissal(self) -> None:
        connection = self.connection_factory()
        persist_source_diagnostics(
            connection, self.search_result(code="partial_results"), "2026-08-08T12:00:00"
        )
        connection.commit()
        history = list_source_diagnostics(connection)
        connection.close()
        self.assertEqual("note", history["items"][0]["level"])
        self.assertEqual("partial_results", history["items"][0]["code"])

    def test_history_is_newest_first_and_retention_is_bounded(self) -> None:
        connection = self.connection_factory()
        with patch.object(source_diagnostics_module, "MAX_HISTORY_RECORDS", 3):
            for index in range(5):
                persist_source_diagnostics(
                    connection,
                    self.search_result(recorded_value=index),
                    f"2026-08-08T12:00:0{index}",
                )
            connection.commit()
        history = list_source_diagnostics(connection, limit=50)
        connection.close()
        self.assertEqual(3, history["count"])
        self.assertEqual(
            ["2026-08-08T12:00:04", "2026-08-08T12:00:03", "2026-08-08T12:00:02"],
            [item["recorded_at"] for item in history["items"]],
        )

    def test_api_retrieves_exports_and_clears_safe_history(self) -> None:
        connection = self.connection_factory()
        persist_source_diagnostics(connection, self.search_result(), "2026-08-08T12:00:00")
        connection.execute(
            "UPDATE source_diagnostics SET counters_json = ?",
            ('{"raw_candidates":4,"leak":"secret-value"}',),
        )
        connection.commit()
        connection.close()
        with patch.object(app_module, "get_db_connection", side_effect=self.connection_factory):
            history_response = app_module.get_source_diagnostic_history()
            export_response = app_module.export_source_diagnostic_history()
            cleared = app_module.clear_source_diagnostic_history()

        history = json.loads(history_response.body)
        exported = json.loads(export_response.body)
        self.assertEqual(1, history["count"])
        self.assertEqual("no-store", history_response.headers["cache-control"])
        self.assertEqual(1, exported["record_count"])
        self.assertIn("attachment;", export_response.headers["content-disposition"])
        exported_text = export_response.body.decode("utf-8")
        self.assertNotIn("private.example", exported_text)
        self.assertNotIn("secret-value", exported_text)
        self.assertEqual(1, cleared["cleared"])

    def test_history_failure_does_not_replace_successful_search_result(self) -> None:
        result = self.search_result()
        result["success"] = True
        connection = unittest.mock.Mock()
        with (
            patch.object(app_module, "run_job_search_and_matching", return_value=result),
            patch.object(app_module, "get_db_connection", return_value=connection),
            patch.object(app_module, "persist_source_diagnostics", side_effect=sqlite3.OperationalError("disk full")),
        ):
            app_module.IS_SEARCHING = True
            app_module.run_search_wrapper("private keywords", "private location")

        self.assertIs(result, app_module.LAST_SEARCH_RESULT)
        self.assertFalse(app_module.IS_SEARCHING)
        connection.rollback.assert_called_once()
        connection.close.assert_called_once()

    def test_history_connection_failure_does_not_replace_successful_search_result(self) -> None:
        result = self.search_result()
        result["success"] = True
        with (
            patch.object(app_module, "run_job_search_and_matching", return_value=result),
            patch.object(app_module, "get_db_connection", side_effect=sqlite3.OperationalError("unavailable")),
        ):
            app_module.IS_SEARCHING = True
            app_module.run_search_wrapper("private keywords", "private location")
        self.assertIs(result, app_module.LAST_SEARCH_RESULT)
        self.assertFalse(app_module.IS_SEARCHING)

    def test_search_without_notices_does_not_open_history_database(self) -> None:
        connection_factory = unittest.mock.Mock()
        result = {"success": True, "provider_alerts": [], "provider_health": {}}
        with (
            patch.object(app_module, "run_job_search_and_matching", return_value=result),
            patch.object(app_module, "get_db_connection", connection_factory),
        ):
            app_module.IS_SEARCHING = True
            app_module.run_search_wrapper("keywords", "location")
        connection_factory.assert_not_called()
        self.assertIs(result, app_module.LAST_SEARCH_RESULT)

    def test_schema_migration_is_idempotent(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "migration.db"
            with patch.object(database_module, "DB_PATH", database_path):
                database_module.init_db()
                database_module.init_db()
            connection = sqlite3.connect(database_path)
            tables = {row[0] for row in connection.execute(
                "SELECT name FROM sqlite_master WHERE type = 'table'"
            )}
            indexes = {row[0] for row in connection.execute(
                "SELECT name FROM sqlite_master WHERE type = 'index'"
            )}
            profile_columns = {row[1] for row in connection.execute("PRAGMA table_info(profile)")}
            application_columns = {row[1] for row in connection.execute("PRAGMA table_info(applications)")}
            base_resume_columns = {row[1] for row in connection.execute("PRAGMA table_info(base_resumes)")}
            base_resume_version_columns = {row[1] for row in connection.execute("PRAGMA table_info(base_resume_versions)")}
            maps_provider = connection.execute("SELECT maps_provider FROM profile WHERE id = 1").fetchone()[0]
            connection.close()
        self.assertIn("source_diagnostics", tables)
        self.assertIn("headquarters_cache", tables)
        self.assertIn("base_resumes", tables)
        self.assertIn("base_resume_versions", tables)
        self.assertIn("idx_source_diagnostics_recorded_at", indexes)
        self.assertIn("idx_base_resume_versions_resume", indexes)
        self.assertIn("maps_provider", profile_columns)
        self.assertIn("active_base_resume_id", profile_columns)
        self.assertIn("evidence_json", base_resume_columns)
        self.assertIn("evidence_json", base_resume_version_columns)
        self.assertEqual("openstreetmap", maps_provider)
        self.assertTrue({
            "headquarters_source", "headquarters_attribution", "base_resume_id",
            "base_resume_name", "base_resume_version",
        }.issubset(application_columns))

    def test_legacy_resume_is_migrated_once_into_named_version_history(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "legacy-resume.db"
            connection = sqlite3.connect(database_path)
            connection.executescript("""
                CREATE TABLE profile (
                    id INTEGER PRIMARY KEY, name TEXT, base_resume_text TEXT, resume_mode TEXT
                );
                INSERT INTO profile VALUES (1, 'Candidate', 'Legacy resume content', 'it');
            """)
            connection.close()
            with patch.object(database_module, "DB_PATH", database_path):
                database_module.init_db()
                database_module.init_db()
            connection = sqlite3.connect(database_path)
            profile = connection.execute(
                "SELECT active_base_resume_id, base_resume_text, resume_mode FROM profile WHERE id = 1"
            ).fetchone()
            resumes = connection.execute(
                "SELECT id, name, content, resume_mode FROM base_resumes"
            ).fetchall()
            versions = connection.execute(
                "SELECT base_resume_id, version_number, content FROM base_resume_versions"
            ).fetchall()
            connection.close()

        self.assertEqual(1, len(resumes))
        self.assertEqual(1, len(versions))
        self.assertEqual("Primary Resume", resumes[0][1])
        self.assertEqual("Legacy resume content", resumes[0][2])
        self.assertEqual("it", resumes[0][3])
        self.assertEqual((resumes[0][0], "Legacy resume content", "it"), profile)
        self.assertEqual((resumes[0][0], 1, "Legacy resume content"), versions[0])

    def test_maps_provider_migration_preserves_google_for_existing_profiles(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "legacy-profile.db"
            connection = sqlite3.connect(database_path)
            connection.executescript("""
                CREATE TABLE profile (id INTEGER PRIMARY KEY, name TEXT);
                INSERT INTO profile VALUES (1, 'Existing Candidate');
            """)
            connection.commit()
            connection.close()
            with patch.object(database_module, "DB_PATH", database_path):
                database_module.init_db()
            connection = sqlite3.connect(database_path)
            provider = connection.execute("SELECT maps_provider FROM profile WHERE id = 1").fetchone()[0]
            connection.close()

        self.assertEqual("google", provider)


class LifecycleSchemaTests(unittest.TestCase):
    def test_lifecycle_columns_and_foreign_keys_are_enabled(self) -> None:
        connection = get_db_connection()
        lifecycle_columns = {
            "created_at", "tailored_at", "form_filled_at", "submitted_at",
            "confirmed_at", "application_method", "submission_evidence", "notes", "follow_up_date", "tailored_resume_text",
            "cover_letter_path", "headquarters_source", "headquarters_attribution",
            "base_resume_id", "base_resume_name", "base_resume_version",
        }
        actual = {row[1] for row in connection.execute("PRAGMA table_info(applications)")}
        self.assertTrue(lifecycle_columns.issubset(actual))
        job_columns = {row[1] for row in connection.execute("PRAGMA table_info(jobs)")}
        self.assertTrue({"last_checked_at", "is_expired", "expiration_reason", "location", "work_arrangement", "employment_type", "compensation", "source"}.issubset(job_columns))
        profile_columns = {row[1] for row in connection.execute("PRAGMA table_info(profile)")}
        self.assertIn("resume_mode", profile_columns)
        self.assertIn("prefer_us_headquarters", profile_columns)
        self.assertIn("active_base_resume_id", profile_columns)
        self.assertTrue({"ai_provider", "ai_model", "openai_api_key", "maps_provider"}.issubset(profile_columns))
        self.assertEqual(1, connection.execute("PRAGMA foreign_keys").fetchone()[0])
        tables = {row[0] for row in connection.execute("SELECT name FROM sqlite_master WHERE type='table'")}
        self.assertTrue({
            "application_status_history", "saved_searches", "job_suppressions",
            "source_diagnostics", "headquarters_cache", "base_resumes",
            "base_resume_versions",
        }.issubset(tables))
        saved_search_columns = {row[1] for row in connection.execute("PRAGMA table_info(saved_searches)")}
        self.assertTrue({"schedule_frequency", "next_alert_at"}.issubset(saved_search_columns))
        connection.close()


class ApplicationInsightsTests(unittest.TestCase):
    def setUp(self) -> None:
        self.connection = sqlite3.connect(":memory:")
        self.connection.row_factory = sqlite3.Row
        self.connection.executescript("""
            CREATE TABLE jobs (
                id INTEGER PRIMARY KEY,
                title TEXT,
                source TEXT,
                location TEXT
            );
            CREATE TABLE applications (
                id INTEGER PRIMARY KEY,
                job_id INTEGER,
                position TEXT,
                date_applied TEXT,
                confirmed_at TEXT,
                created_at TEXT,
                application_method TEXT,
                base_resume_name TEXT,
                base_resume_version INTEGER,
                status TEXT
            );
            CREATE TABLE application_status_history (
                id INTEGER PRIMARY KEY,
                job_id INTEGER,
                to_status TEXT,
                changed_at TEXT,
                undone_at TEXT
            );
        """)
        self.connection.executemany(
            "INSERT INTO jobs (id, title, source, location) VALUES (?, ?, ?, ?)",
            (
                (1, "Data Architect", "lever", "Remote"),
                (2, "Chief Data Officer", "greenhouse", "Dayton, OH"),
                (3, "VP Data", "greenhouse", "Dayton, OH"),
                (4, "Security Architect", "lever", "Remote"),
                (5, "Tailored Only", "ashby", "Chicago, IL"),
            ),
        )
        self.connection.executemany("""
            INSERT INTO applications (
                id, job_id, position, date_applied, confirmed_at, created_at,
                application_method, base_resume_name, base_resume_version, status
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            (1, 1, "Data Architect", "2026-08-01", "2026-08-01T09:00:00", "2026-08-01", "manual:company_site", "Architecture", 3, "applied"),
            (2, 2, "Chief Data Officer", "2026-08-01", "2026-08-01T09:00:00", "2026-08-01", "manual:referral", "Executive", 2, "rejected"),
            (3, 3, "VP Data", "2026-08-02", "2026-08-02T09:00:00", "2026-08-02", "manual:referral", "Executive", 2, "offer"),
            (4, 4, "Security Architect", "2026-08-02", "2026-08-02T09:00:00", "2026-08-02", "job_board", None, None, "applied"),
            (5, 5, "Tailored Only", None, None, "2026-08-02", None, "General", 1, "tailored"),
        ))
        self.connection.executemany("""
            INSERT INTO application_status_history (id, job_id, to_status, changed_at, undone_at)
            VALUES (?, ?, ?, ?, ?)
        """, (
            (1, 2, "interview", "2026-08-03T10:00:00", None),
            (2, 2, "rejected", "2026-08-05T10:00:00", None),
            (3, 3, "offer", "2026-08-05T10:00:00", None),
            (4, 4, "interview", "2026-08-04T10:00:00", "2026-08-04T11:00:00"),
        ))

    def tearDown(self) -> None:
        self.connection.close()

    def test_summarizes_confirmed_outcomes_and_active_history(self) -> None:
        result = build_application_insights(self.connection)

        self.assertEqual({
            "applications": 4,
            "responses": 2,
            "positive_responses": 2,
            "interviews": 1,
            "offers": 1,
            "rejections": 1,
            "pending": 2,
            "response_rate": 50.0,
            "positive_response_rate": 50.0,
            "average_response_days": 2.5,
        }, result["summary"])
        self.assertEqual(
            "An interview, offer, or rejection recorded in the current lifecycle or its active history.",
            result["definitions"]["response"],
        )

    def test_groups_by_each_requested_dimension(self) -> None:
        result = build_application_insights(self.connection)

        by_source = {item["label"]: item for item in result["groups"]["source"]}
        self.assertEqual(2, by_source["Greenhouse"]["applications"])
        self.assertEqual(100.0, by_source["Greenhouse"]["response_rate"])
        self.assertEqual(0.0, by_source["Lever"]["response_rate"])
        by_location = {item["label"]: item for item in result["groups"]["location"]}
        self.assertEqual(2, by_location["Dayton, OH"]["responses"])
        by_resume = {item["label"]: item for item in result["groups"]["resume"]}
        self.assertEqual(2, by_resume["Executive · v2"]["applications"])
        self.assertIn("Unattributed resume", by_resume)
        by_method = {item["label"]: item for item in result["groups"]["method"]}
        self.assertEqual(2, by_method["Referral"]["applications"])
        by_role = {item["label"]: item for item in result["groups"]["role"]}
        self.assertIn("Chief Data Officer", by_role)
        self.assertNotIn("Tailored Only", by_role)

    def test_normalizes_common_source_hostnames_for_display(self) -> None:
        self.connection.execute("UPDATE jobs SET source = 'www.linkedin.com' WHERE id = 1")

        result = build_application_insights(self.connection)

        labels = {item["label"] for item in result["groups"]["source"]}
        self.assertIn("LinkedIn", labels)
        self.assertNotIn("Www.Linkedin.Com", labels)


class ManualLifecycleTests(unittest.TestCase):
    def setUp(self) -> None:
        self.connection = sqlite3.connect(":memory:")
        self.connection.row_factory = sqlite3.Row
        self.connection.executescript("""
            CREATE TABLE jobs (
                id INTEGER PRIMARY KEY, company TEXT, title TEXT, status TEXT
            );
            CREATE TABLE applications (
                id INTEGER PRIMARY KEY, job_id INTEGER UNIQUE, company TEXT, position TEXT,
                date_applied TEXT, status TEXT, created_at TEXT, submitted_at TEXT,
                confirmed_at TEXT, application_method TEXT, submission_evidence TEXT,
                notes TEXT, follow_up_date TEXT, tailored_resume_path TEXT, cover_letter TEXT
            );
            CREATE TABLE application_status_history (
                id INTEGER PRIMARY KEY, job_id INTEGER, from_status TEXT, to_status TEXT,
                changed_at TEXT, source TEXT, notes TEXT, undone_at TEXT
            );
            INSERT INTO jobs VALUES (1, 'Example Co', 'Data Director', 'matched');
        """)

    def tearDown(self) -> None:
        self.connection.close()

    def test_manual_application_records_date_method_and_evidence(self) -> None:
        result = update_lifecycle(
            self.connection, 1, "applied", applied_on=date(2026, 8, 3),
            method="referral", notes="Referred by Alex", source="manual",
        )
        self.assertEqual("applied", result["status"])
        application = self.connection.execute("SELECT * FROM applications WHERE job_id = 1").fetchone()
        self.assertEqual("2026-08-03", application["date_applied"])
        self.assertEqual("manual:referral", application["application_method"])
        self.assertIn('"confirmed_by_user": true', application["submission_evidence"])
        self.assertEqual("applied", self.connection.execute("SELECT status FROM jobs WHERE id=1").fetchone()[0])

    def test_latest_manual_change_can_be_undone(self) -> None:
        update_lifecycle(self.connection, 1, "applied", applied_on=date(2026, 8, 3), source="manual")
        result = undo_latest_lifecycle_change(self.connection, 1)
        self.assertEqual("matched", result["status"])
        self.assertEqual("matched", self.connection.execute("SELECT status FROM jobs WHERE id=1").fetchone()[0])
        self.assertIsNone(self.connection.execute("SELECT id FROM applications WHERE job_id=1").fetchone())


class BulkCleanupTests(unittest.TestCase):
    def setUp(self) -> None:
        self.connection = sqlite3.connect(":memory:")
        self.connection.row_factory = sqlite3.Row
        self.connection.executescript("""
            CREATE TABLE jobs (
                id INTEGER PRIMARY KEY,
                company TEXT,
                title TEXT,
                url TEXT,
                date_found TEXT,
                match_score INTEGER,
                status TEXT,
                archived_at TEXT,
                archived_from_status TEXT
            );
            CREATE TABLE applications (id INTEGER PRIMARY KEY, job_id INTEGER);
            CREATE TABLE job_suppressions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                url_fingerprint TEXT NOT NULL UNIQUE, hostname TEXT NOT NULL,
                company TEXT NOT NULL DEFAULT '', title TEXT NOT NULL DEFAULT '',
                deleted_at TEXT NOT NULL, deletion_source TEXT NOT NULL
            );
            INSERT INTO jobs VALUES (1, 'A', 'Untouched', 'https://jobs.example.test/1', '2026-01-01', 90, 'matched', NULL, NULL);
            INSERT INTO jobs VALUES (2, 'B', 'Has History', 'https://jobs.example.test/2', '2026-01-01', 80, 'matched', NULL, NULL);
            INSERT INTO jobs VALUES (3, 'C', 'Tailored', 'https://jobs.example.test/3', '2026-01-01', 70, 'tailored', NULL, NULL);
            INSERT INTO jobs VALUES (4, 'D', 'Archived', 'https://jobs.example.test/4', '2026-01-01', 60, 'archived', '2026-01-02', 'matched');
            INSERT INTO applications VALUES (1, 2);
        """)

    def tearDown(self) -> None:
        self.connection.close()

    def test_preview_protects_any_job_with_history_or_progress(self) -> None:
        preview = cleanup_preview(self.connection)
        self.assertEqual(1, preview["actions"]["archive"]["count"])
        self.assertEqual(2, preview["actions"]["delete"]["count"])
        self.assertEqual(1, preview["actions"]["restore"]["count"])
        self.assertEqual(2, preview["protected_count"])

    def test_preview_token_prevents_changed_set_from_being_mutated(self) -> None:
        preview = cleanup_preview(self.connection)
        token = preview["actions"]["archive"]["preview_token"]
        self.connection.execute(
            "INSERT INTO jobs VALUES (5, 'E', 'New', 'https://jobs.example.test/5', '2026-01-03', 50, 'matched', NULL, NULL)"
        )
        with self.assertRaisesRegex(ValueError, "changed after preview"):
            apply_cleanup(self.connection, "archive", token, "2026-01-04T00:00:00")
        status = self.connection.execute("SELECT status FROM jobs WHERE id = 1").fetchone()[0]
        self.assertEqual("matched", status)

    def test_archive_and_restore_are_reversible(self) -> None:
        preview = cleanup_preview(self.connection)
        affected = apply_cleanup(
            self.connection, "archive", preview["actions"]["archive"]["preview_token"], "2026-01-04T00:00:00"
        )
        self.assertEqual(1, affected)
        self.assertEqual("archived", self.connection.execute("SELECT status FROM jobs WHERE id = 1").fetchone()[0])

        preview = cleanup_preview(self.connection)
        affected = apply_cleanup(
            self.connection, "restore", preview["actions"]["restore"]["preview_token"], "2026-01-04T00:00:00"
        )
        self.assertEqual(2, affected)
        restored = self.connection.execute("SELECT status FROM jobs WHERE id IN (1, 4) ORDER BY id").fetchall()
        self.assertEqual(["matched", "matched"], [row[0] for row in restored])

    def test_delete_removes_only_previewed_untouched_jobs(self) -> None:
        preview = cleanup_preview(self.connection)
        affected = apply_cleanup(
            self.connection, "delete", preview["actions"]["delete"]["preview_token"], "2026-01-04T00:00:00"
        )
        self.assertEqual(2, affected)
        remaining = self.connection.execute("SELECT id FROM jobs ORDER BY id").fetchall()
        self.assertEqual([2, 3], [row[0] for row in remaining])
        suppressions = self.connection.execute(
            "SELECT hostname, deletion_source FROM job_suppressions ORDER BY id"
        ).fetchall()
        self.assertEqual(
            [("jobs.example.test", "bulk_cleanup"), ("jobs.example.test", "bulk_cleanup")],
            [tuple(row) for row in suppressions],
        )


if __name__ == "__main__":
    unittest.main()
